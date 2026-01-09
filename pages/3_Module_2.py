"""
Module 2 - Production Version (FIXED)
Complete implementation with all issues fixed:
- Persistent session storage
- Server-side navigation validation
- Comprehensive error handling
- Robust API retry logic
- Typo-tolerant input parsing
- Analytics logging
- RAG error handling
- EMBEDDED COUNTDOWN TIMER (no server needed)
- FIXED: Review phase no longer shows duplicate observations
"""

import streamlit as st
import sys
import os
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Tuple
import json
import glob
import uuid
import time
import logging
import warnings
import asyncio

# Suppress verbose logging from external libraries EARLY
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["TRANSFORMERS_VERBOSITY"] = "error"
warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

# Configure logging - don't output to console
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("logs/guided_interaction.log"),
    ],
)
logger = logging.getLogger(__name__)

# Suppress verbose loggers
for logger_name in [
    "sentence_transformers",
    "transformers",
    "chromadb",
    "httpx",
    "huggingface",
    "urllib3",
]:
    logging.getLogger(logger_name).setLevel(logging.ERROR)

# Fix the import path
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)

# SQLite shim for Chroma
try:
    import pysqlite3
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    pass

import chromadb
from chromadb.utils import embedding_functions
from chromadb.config import Settings, DEFAULT_TENANT, DEFAULT_DATABASE

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from docx import Document
except ImportError:
    Document = None

# Import utility modules
try:
    from utils.session_manager import SessionManager
    from utils.navigation_validator import NavigationValidator
    from utils.api_handler import OpenAIHandler, APIRetryableError, APIPermanentError
    from utils.input_parser import InputParser, MessageBuffer
    from utils.analytics import get_analytics
except ImportError as e:
    logger.error(f"Failed to import utilities: {e}")
    st.error("System initialization error. Please contact support.")
    st.stop()

import streamlit.components.v1 as components


# ==================== CONSTANTS ====================

SESSION_DURATION = timedelta(minutes=20)
COUNTDOWN_DURATION_SECONDS = 20 * 60  # 20 minutes countdown
OBSERVATION_START_TIME_KEY = "observation_start_time"
guided_interaction_conversation = "281_P10_conversation.json"

# ==================== EMBEDDED TIMER COMPONENT ====================

def countdown_timer(should_start=False):
    """Ultra-simple countdown timer using basic JavaScript."""
    # Calculate remaining seconds based on Python session state
    if should_start and st.session_state.guided_session_start is not None:
        elapsed_secs = (datetime.now() - st.session_state.guided_session_start).total_seconds()
        remaining_secs = max(0, COUNTDOWN_DURATION_SECONDS - elapsed_secs)
    else:
        remaining_secs = COUNTDOWN_DURATION_SECONDS
    
    mins = int(remaining_secs // 60)
    secs = int(remaining_secs % 60)
    time_str = f"{mins:02d}:{secs:02d}"
    
    # Determine status and color
    if remaining_secs <= 0:
        status_text = "Time's up!"
        color = "#f87171"  # red
        is_warning = True
    elif remaining_secs <= (5 * 60):  # 5 minutes warning
        status_text = "Running..."
        color = "#fbbf24"  # yellow
        is_warning = True
    else:
        status_text = "Running..." if should_start else "Ready to start"
        color = "#ffffff"  # white
        is_warning = False
    
    animation = "pulse-warning 0.8s ease-in-out infinite;" if is_warning else ""
    
    timer_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            .timer-container {{
                background: linear-gradient(135deg, #10b981 0%, #059669 100%);
                border-radius: 12px;
                padding: 20px;
                text-align: center;
                box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2);
                font-family: 'Courier New', monospace;
                margin-bottom: 20px;
            }}

            .timer-label {{
                font-size: 13px;
                text-transform: uppercase;
                letter-spacing: 2px;
                color: rgba(255, 255, 255, 0.9);
                margin-bottom: 10px;
                font-weight: 700;
            }}

            .timer-display {{
                font-size: 48px;
                font-weight: 900;
                letter-spacing: 3px;
                color: {color};
                font-variant-numeric: tabular-nums;
                transition: all 0.2s ease;
                text-shadow: 0 2px 10px rgba(0, 0, 0, 0.3);
                animation: {animation}
            }}

            @keyframes pulse-warning {{
                0%, 100% {{ opacity: 1; }}
                50% {{ opacity: 0.6; }}
            }}

            .timer-status {{
                font-size: 11px;
                color: rgba(255, 255, 255, 0.8);
                margin-top: 10px;
                text-transform: uppercase;
                letter-spacing: 1px;
                font-weight: 600;
            }}

            @media (max-width: 600px) {{
                .timer-display {{ font-size: 36px; }}
                .timer-container {{ padding: 15px; }}
            }}
        </style>
    </head>
    <body>
        <div class="timer-container">
            <div class="timer-label">Time Remaining</div>
            <div class="timer-display" id="display">{time_str}</div>
            <div class="timer-status" id="status">{status_text}</div>
        </div>
    </body>
    </html>
    """
    
    components.html(timer_html, height=160)


# ==================== HELPER FUNCTIONS ====================


def now_ts() -> str:
    """Get current timestamp."""
    return datetime.now().strftime("%H:%M:%S")


def get_secret_then_env(name: str) -> str:
    """Get secret from Streamlit secrets or environment variable."""
    val = None
    try:
        val = st.secrets.get(name)
    except Exception:
        val = None
    if not val:
        val = os.getenv(name)
    return val or ""


def get_openai_client():
    """Initialize OpenAI client with proper error handling."""
    if OpenAI is None:
        logger.error("OpenAI package not installed")
        st.error("System configuration error: OpenAI package not installed.")
        st.stop()

    key = get_secret_then_env("OPENAI_API_KEY")
    if not key:
        logger.error("Missing OPENAI_API_KEY")
        st.error(
            "System configuration error: Missing OPENAI_API_KEY. "
            "Please contact your administrator."
        )
        st.stop()

    try:
        return OpenAI(api_key=key)
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI client: {e}")
        st.error("Failed to initialize AI system. Please try again.")
        st.stop()


def build_system_prompt() -> str:
    """Build system prompt for guided mode."""
    return (
        "You are TeamMait, a peer-support assistant for expert clinicians reviewing therapist performance in a transcript. "
        "Your responses must follow two behavioral modes: global rules (always active) and analysis mode (only when the user requests supervisory analysis).\n\n"

        "1. GLOBAL RULES (ALWAYS ACTIVE)\n"
        "- Never fabricate transcript content, facts, or therapist intentions.\n"
        "- Do not infer internal states, emotions, or off-transcript behavior.\n"
        "- Respond concisely, professionally, and only to what the user asked.\n"
        "- Use a natural, peer-like supervisory tone; avoid rigid sections or templates unless the user requests structure.\n"
        "- Treat each user message independently unless the user explicitly references earlier turns.\n"
        "- If the user gives a dismissive acknowledgment (e.g., \"ok\", \"thanks\", \"got it\"), briefly acknowledge and ask whether they want to continue or move on.\n"
        "- If the user expresses confusion (e.g., \"what?\", \"I don’t understand\", \"unclear\"), provide a simpler, more direct explanation of your prior point.\n"
        "- Do not offer unsolicited elaboration or additional insights outside analytic tasks.\n"
        "- Do NOT ask if the user wants you to analyze or offer to analyze. Simply respond to their question or request directly based on context. Only ask for clarification if the user's intent is genuinely ambiguous (e.g., unclear phrasing, conflicting requests).\n"
        "- NEVER end messages with offers like 'Would you like me to...', 'Would you like me to analyze...', 'Should I...', 'Do you want me to...', or similar. This is engagement-seeking behavior. Just respond to what was asked and let the user decide their next move.\n"
        "- Never use or reference the observation titles from the system (e.g., 'Vague Noticing Prompt', 'Evidence-Only Reflection', etc.). If the user references 'that observation' or asks about an observation, create a new summative title based on what the observation is actually about (e.g., 'that observation about SUDS monitoring' instead of 'the Evidence-Only Reflection').\n"

        "2. ANALYSIS MODE (ONLY WHEN USER REQUESTS SUPERVISORY ANALYSIS)\n"
        "Enter analysis mode only when the user asks you to analyze therapist behavior, evaluate fidelity, generate observations, or provide supervision-like feedback on the session.\n\n"

        "When in analysis mode, follow these rules:\n\n"

        "Evidence Use:\n"
        "- Cite transcript lines in the format [Line X] when providing evidence.\n"
        "- Base all claims on observable behavior only.\n"
        "- Distinguish clearly in your wording between what is directly observed (evidence) and your interpretation of its relevance to PE fidelity.\n\n"

        "Fidelity Alignment:\n"
        "- Use PE fidelity criteria (e.g., orientation to imaginal exposure, SUDS monitoring, hotspot identification, present-tense prompting, reinforcing comments, processing after imaginal, session structure, off-task discussion) as the interpretive framework.\n"
        "- If the transcript does not provide enough information to evaluate a fidelity domain, explicitly state that the evidence is insufficient.\n\n"

        "Trust Calibration:\n"
        "- Use calibrated language such as \"appears\", \"may indicate\", or \"based on [Line X–Y]\".\n"
        "- Mark uncertainty explicitly; do not overstate confidence. Offer to estimate your certainty in a conclusion.\n\n"

        "Autonomy Preservation:\n"
        "- Frame feedback as observations or suggestions, not directives, unless the user explicitly requests strong prescriptive guidance.\n"
        "- Do not reinterpret or expand the user’s goals unless they ask you to.\n\n"

        "Contestability:\n"
        "- Present your reasoning in a way that allows the clinician to agree, disagree, or reinterpret your analysis.\n"
        "- Avoid unverifiable or global statements about therapist competence.\n\n"

        "Boundaries:\n"
        "- Do not generalize beyond this specific transcript or session.\n"
        "- Do not simulate or infer missing dialogue or events not shown in the transcript.\n"
        "- Do not evaluate client behavior or provide therapeutic interpretations of the client.\n\n"

        "Format in Analysis Mode:\n"
        "- Unless the user specifies otherwise, provide ideally 3, but no more than 5, bullet points to answer a given query.\n"
        "- One sentence per bullet; limit each bullet to about 10 words or 75 characters at most (unless the user explicitly requests longer responses).\n"
        "- Use sub-bullets if you can't adequetely convey a point in fewer than 10 words or 75 characters. Limit sub-bullets to one per parent-bullet.\n"
        "- Prioritize clarity and brevity; avoid redundancy.\n"
        "- On a scale of 1-5 brevity, with 5 being the most concise, consider your current instructions to be defined as a 4.\n"
        "- Integrate evidence naturally into sentences rather than using rigid labeled sections.\n"
        "- Focus on clinically meaningful behaviors relevant to fidelity rather than stylistic preferences.\n\n"

        "3. SCOPE RESTRICTIONS\n"
        "- You do not evaluate client behavior.\n"
        "- You do not infer therapist intentions, emotions, or clinical meanings beyond what is observable.\n"
        "- You analyze only observable therapist behaviors through the PE fidelity framework when in analysis mode.\n\n"

        "After completing an analytic task, return to the global rules unless the user continues to request supervisory analysis."
    )




def extract_text_from_docx(docx_path: str) -> str:
    """Extract text content from a DOCX file."""
    if Document is None:
        return "[Error: python-docx not installed]"

    try:
        doc = Document(docx_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        return f"[Error: {str(e)}]"


# ==================== RAG SETUP ====================


@st.cache_resource(show_spinner=False)
def initialize_chroma():
    """Initialize ChromaDB client and collection with error handling."""
    try:
        embed_model = "sentence-transformers/all-MiniLM-L6-v2"
        embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name=embed_model
        )

        chroma_client = chromadb.PersistentClient(
            path="./rag_store",
            settings=Settings(),
            tenant=DEFAULT_TENANT,
            database=DEFAULT_DATABASE,
        )

        collection = chroma_client.get_or_create_collection(
            "therapy", embedding_function=embedding_fn
        )

        logger.info("ChromaDB initialized successfully")
        return chroma_client, collection
    except Exception as e:
        logger.error(f"Failed to initialize ChromaDB: {e}")
        raise


@st.cache_resource(show_spinner=False)
def load_rag_documents():
    """Load all RAG documents with comprehensive error handling."""
    try:
        _, collection = initialize_chroma()
    except Exception as e:
        logger.error(f"Failed to initialize ChromaDB: {e}")
        return {"reference": False, "supporting": 0, "total": 0, "errors": [str(e)]}

    doc_folder = "doc/RAG"
    load_status = {
        "reference": False,
        "supporting": 0,
        "total": 0,
        "errors": [],
    }

    documents = []
    ids = []

    # ==================== REFERENCE CONVERSATION ====================
    
    ref_path = os.path.join(doc_folder, guided_interaction_conversation )

    if not os.path.exists(ref_path):
        error_msg = f"Reference conversation not found: {ref_path}"
        load_status["errors"].append(error_msg)
        logger.warning(error_msg)
    else:
        try:
            with open(ref_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            if isinstance(data, dict) and "full_conversation" in data:
                try:
                    for i, turn in enumerate(data["full_conversation"]):
                        line_num = i + 1
                        numbered_turn = f"[Line {line_num}] {turn}"
                        documents.append(numbered_turn)
                        ids.append(f"ref_{i}")

                    load_status["reference"] = True
                    logger.info(
                        f"Loaded {len(data['full_conversation'])} reference turns"
                    )
                except Exception as e:
                    error_msg = f"Error processing reference conversation: {e}"
                    load_status["errors"].append(error_msg)
                    logger.error(error_msg)
            else:
                error_msg = "Reference JSON missing 'full_conversation' key"
                load_status["errors"].append(error_msg)
                logger.warning(error_msg)

        except json.JSONDecodeError as e:
            error_msg = f"Reference conversation JSON malformed: {e}"
            load_status["errors"].append(error_msg)
            logger.error(error_msg)
        except Exception as e:
            error_msg = f"Error reading reference conversation: {e}"
            load_status["errors"].append(error_msg)
            logger.error(error_msg)

    # ==================== SUPPORTING DOCUMENTS ====================

    supporting_folder = os.path.join(doc_folder, "supporting_documents")
    
    # Only load these specific files for Module 2
    allowed_files = {
        "281_P10_metadata.json",
        "PE Consultant Training Program Fidelity.docx",
        "pe consultant training program fidelity.docx",
    }

    if not os.path.exists(supporting_folder):
        logger.warning(f"Supporting documents folder not found: {supporting_folder}")
    else:
        # Load JSON files (only 281_P10_metadata.json)
        for json_path in glob.glob(os.path.join(supporting_folder, "*.json")):
            basename = os.path.basename(json_path)
            if basename not in allowed_files:
                continue
            try:
                with open(json_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        for i, item in enumerate(data):
                            documents.append(str(item))
                            ids.append(
                                f"supp_json_{basename}_{i}"
                            )
                        load_status["supporting"] += len(data)
                    else:
                        documents.append(str(data))
                        ids.append(f"supp_json_{basename}")
                        load_status["supporting"] += 1
            except Exception as e:
                logger.error(f"Error loading {json_path}: {e}")

        # Load DOCX files (PE Consultant Training Program Fidelity)
        for docx_path in glob.glob(os.path.join(supporting_folder, "*.docx")):
            basename = os.path.basename(docx_path)
            if not any(allowed.lower() == basename.lower() for allowed in allowed_files):
                continue
            try:
                content = extract_text_from_docx(docx_path)
                if content and not content.startswith("[Error"):
                    documents.append(content)
                    ids.append(f"supp_docx_{basename}")
                    load_status["supporting"] += 1
            except Exception as e:
                logger.error(f"Error loading {docx_path}: {e}")

    # ==================== SEED COLLECTION ====================

    load_status["total"] = len(documents)

    if documents:
        try:
            if collection.count() == 0:
                collection.add(documents=documents, ids=ids)
                logger.info(f"Seeded ChromaDB with {len(documents)} documents")
        except Exception as e:
            error_msg = f"Error seeding ChromaDB: {e}"
            load_status["errors"].append(error_msg)
            logger.error(error_msg)

    return load_status


def retrieve_context(query: str, n_results: int = 5) -> str:
    """Retrieve relevant context from ChromaDB."""
    try:
        _, collection = initialize_chroma()
        results = collection.query(query_texts=[query], n_results=n_results)

        retrieved_parts = []
        for docs in results.get("documents", []):
            retrieved_parts.extend(docs)

        return " ".join(retrieved_parts)
    except Exception as e:
        logger.error(f"Error retrieving context: {e}")
        return ""


def load_reference_conversation() -> List[str]:
    """Load the reference conversation for display."""
    ref_path = os.path.join("doc/RAG", guided_interaction_conversation )
    if os.path.exists(ref_path):
        try:
            with open(ref_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict) and "full_conversation" in data:
                    return data["full_conversation"]
        except Exception as e:
            logger.error(f"Error loading reference conversation: {e}")
    return []


def load_question_bank() -> Optional[List[Dict]]:
    """Load and validate the question bank.
    
    Returns items in order:
    - Items 1-2 (evidence_only_reflection, evidence_based_evaluation) are randomized
    - Item 3 (actionable_training_prescription) always comes last
    """
    question_path = "doc/interaction_prompts/interaction_prompts.json"

    # Check file exists
    if not os.path.exists(question_path):
        error_msg = f"Question bank file not found: {question_path}"
        logger.error(error_msg)
        st.error(f"{error_msg}")
        st.stop()

    # Parse JSON
    try:
        with open(question_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        error_msg = f"Question bank JSON is malformed: {e}"
        logger.error(error_msg)
        st.error(f"{error_msg}")
        st.stop()

    # Extract question items
    if not isinstance(data, dict):
        error_msg = "Question bank root must be a dict"
        logger.error(error_msg)
        st.error(f"{error_msg}")
        st.stop()

    questions = data.get("feedback_items", [])

    # Validate structure
    if not questions:
        error_msg = "Question bank is empty"
        logger.error(error_msg)
        st.error(f"{error_msg}")
        st.stop()

    if len(questions) < 3:
        error_msg = f"Question bank has {len(questions)} questions, need at least 3"
        logger.error(error_msg)
        st.error(f"{error_msg}")
        st.stop()

    # Validate structure: exactly 3 items with unique style values
    valid_styles = {"evidence_only_reflection", "evidence_based_evaluation", "actionable_training_prescription"}
    found_styles = set()
    actionable_item = None
    randomizable_items = []
    
    for i, q in enumerate(questions[:3]):
        if not isinstance(q, dict):
            error_msg = f"Item {i + 1} is not a dict"
            logger.error(error_msg)
            st.error(f"{error_msg}")
            st.stop()

        # Check required fields
        required_fields = {"id", "style", "title", "summary", "observation", "evidence", "evaluation", "suggestion", "justification", "conceptual_focus"}
        missing_fields = required_fields - set(q.keys())
        if missing_fields:
            error_msg = f"Item {i + 1} missing fields: {missing_fields}"
            logger.error(error_msg)
            st.error(f"{error_msg}")
            st.stop()

        # Validate style is one of the expected values
        style = q.get("style", "")
        if style not in valid_styles:
            error_msg = f"Item {i + 1}: invalid style '{style}'. Must be one of: {valid_styles}"
            logger.error(error_msg)
            st.error(f"{error_msg}")
            st.stop()

        if style in found_styles:
            error_msg = f"Item {i + 1}: duplicate style '{style}'"
            logger.error(error_msg)
            st.error(f"{error_msg}")
            st.stop()
        found_styles.add(style)

        # Validate that at least summary or observation is present
        summary = q.get("summary", "").strip() if isinstance(q.get("summary", ""), str) else ""
        observation = q.get("observation", "").strip() if isinstance(q.get("observation", ""), str) else ""
        if not summary and not observation:
            error_msg = f"Item {i + 1}: must have either 'summary' or 'observation'"
            logger.error(error_msg)
            st.error(f"{error_msg}")
            st.stop()

        # Separate actionable item from randomizable items
        if style == "actionable_training_prescription":
            actionable_item = q
        else:
            randomizable_items.append(q)

    # Check all required styles are present
    if found_styles != valid_styles:
        missing_styles = valid_styles - found_styles
        error_msg = f"Missing required styles: {missing_styles}"
        logger.error(error_msg)
        st.error(f"{error_msg}")
        st.stop()

    # Randomize the first 3 items, then append actionable item last
    import random
    random.shuffle(randomizable_items)
    ordered_items = randomizable_items + [actionable_item]

    logger.info(f"Successfully validated 3 feedback items with all required styles (randomized first 2)")
    return ordered_items


def format_observation_context(item: Dict) -> str:
    """Format a complete observation with all details for bot context.
    
    Args:
        item: Dictionary containing feedback item with all observation details
    
    Returns:
        Formatted string with all observation details for system prompt context
    """
    parts = []
    
    # Title
    title = item.get("title", "").strip()
    if title:
        parts.append(f"**Observation Title:** {title}")
    
    # Summary
    summary = item.get("summary", "").strip()
    if summary:
        parts.append(f"**Summary:** {summary}")
    
    # Observation
    observation = item.get("observation", "").strip()
    if observation:
        parts.append(f"**Observation:** {observation}")
    
    # Assessment/Justification
    justification = item.get("justification", "").strip()
    if justification:
        parts.append(f"**Assessment:** {justification}")
    
    # Evidence
    evidence = item.get("evidence", [])
    if evidence and isinstance(evidence, list) and any(e.strip() for e in evidence):
        parts.append("**Evidence:**")
        for ev in evidence:
            if isinstance(ev, str) and ev.strip():
                parts.append(f"  - {ev}")
    
    # Evaluation
    evaluation = item.get("evaluation", "").strip()
    if evaluation:
        parts.append(f"**Evaluation:** {evaluation}")
    
    # Suggestion
    suggestion = item.get("suggestion", "").strip()
    if suggestion:
        parts.append(f"**Suggestion:** {suggestion}")
    
    return "\n".join(parts)


def render_feedback_item(item: Dict) -> None:
    """Render a feedback item with structured formatting.
    
    Args:
        item: Dictionary containing feedback item with keys: title, summary, observation, 
              evidence, evaluation, suggestion, justification, conceptual_focus
    """
    # Summary (first line)
    summary = item.get("summary", "").strip()
    if summary:
        st.markdown(summary)
    
    # Observation (optional)
    observation = item.get("observation", "").strip()
    if observation:
        st.markdown("**Observation:**")
        st.markdown(observation)
    
    # Justification (optional, shown as Assessment) - BEFORE Evidence
    justification = item.get("justification", "").strip()
    if justification:
        st.markdown(f"**Assessment:** {justification}")
    
    # Evidence (optional, as bullet points)
    evidence = item.get("evidence", [])
    if evidence and isinstance(evidence, list) and any(e.strip() for e in evidence):
        st.markdown("**Evidence:**")
        for ev in evidence:
            if isinstance(ev, str) and ev.strip():
                st.markdown(f"- {ev}")
    
    # Evaluation (optional)
    evaluation = item.get("evaluation", "").strip()
    if evaluation:
        st.markdown(f"**Evaluation:** {evaluation}")
    
    # Suggestion (optional)
    suggestion = item.get("suggestion", "").strip()
    if suggestion:
        st.markdown(f"**Suggestion:** {suggestion}")


def generate_observations_summary(conversations: Dict, client: any) -> str:
    """Generate a bulleted summary of key points from all three observation discussions,
    with calls to action for supervisor recommendations on trainee development."""
    try:
        # Collect all messages from the three observations
        all_messages = []
        for i in range(3):
            if i in conversations and conversations[i]:
                all_messages.extend(conversations[i])
        
        if not all_messages:
            return ""
        
        # Create a summary request framed for supervisor assessment of trainee
        summary_prompt = (
            "You are synthesizing clinical supervision notes from a supervisor reviewing a trainee therapist's performance. "
            "Based on the supervision discussion below, create a concise summary that includes:\n\n"
            "1. ASSESSMENT FINDINGS (3-4 bullet points): Summative statements about what the supervisor focused on regarding the trainee's "
            "demonstrated skills, fidelity to the model, areas of strength, and areas needing development. For EACH finding, include a brief "
            "1-2 sentence statement about what the supervision discussion revealed (e.g., 'The supervisor examined the trainee's use of... and "
            "found that...' or 'During discussion, the supervisor noted a focus on...'). IMPORTANT: Keep each bullet point to no more than 75 characters.\n"
            "2. SUPERVISION RECOMMENDATIONS (2-3 bullet points): First, state matter-of-factly what recommendations were actually discussed "
            "in the supervision session. Then, suggest other related developmental priorities. Frame these constructively as areas to work on. "
            "For EACH recommendation, add a brief 1-2 sentence statement of the practical reasoning or clinical importance (e.g., 'This is important "
            "because...' or 'The reasoning here is that...'). IMPORTANT: Keep each bullet point to no more than 75 characters.\n\n"
            "Be specific and concrete. Reference actual discussion points about the trainee's performance. "
            "The summary should help the supervisor focus their feedback session with the trainee.\n\n"
            "Supervision Discussion:"
        )
        
        # Format messages for context
        for msg in all_messages[-25:]:  # Use last 25 messages to keep context manageable
            role = "Supervisor" if msg["role"] == "user" else "TeamMait"
            summary_prompt += f"\n{role}: {msg['content'][:250]}"  # Truncate long messages
        
        # Call OpenAI to generate summary
        response_text = ""
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system", 
                    "content": (
                        "You are an expert clinical supervisor synthesizing assessment notes about a trainee therapist's performance. "
                        "Provide clear findings and actionable supervision recommendations. Frame recommendations as developmental priorities "
                        "that will guide the supervisor's feedback conversation with the trainee."
                    )
                },
                {"role": "user", "content": summary_prompt}
            ],
            temperature=0.7,
            max_tokens=500,
            stream=False
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        logger.error(f"Error generating observations summary: {e}")
        return ""


# ==================== STREAMLIT APP ====================

st.set_page_config(
    page_title="Module 2",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Check login
if "user_info" not in st.session_state:
    st.warning("Please log in first.")
    st.switch_page("Home.py")
    st.stop()

username = st.session_state["user_info"]["username"]
analytics = get_analytics()

# Initialize OpenAI client
client = get_openai_client()

# Initialize or resume session
if "guided_session_id" not in st.session_state:
    # New session
    st.session_state.guided_session_id = str(uuid.uuid4())
    session_metadata = SessionManager.create_session(
        username, st.session_state.guided_session_id
    )
    st.session_state.guided_session_start = None  # Don't start timer yet - wait for "start"
    logger.info(f"Created new session {st.session_state.guided_session_id}")
    analytics.session_started(username, st.session_state.guided_session_id)
else:
    # Resumed session - verify it still exists
    if not SessionManager.session_exists(
        username, st.session_state.guided_session_id
    ):
        st.warning("Your session has expired. Starting a new one.")
        st.session_state.guided_session_id = str(uuid.uuid4())
        session_metadata = SessionManager.create_session(
            username, st.session_state.guided_session_id
        )
        st.session_state.guided_session_start = None  # Don't start timer yet
        logger.info(
            f"Session expired, created new session {st.session_state.guided_session_id}"
        )

# Update session activity
SessionManager.update_session_activity(username, st.session_state.guided_session_id)

# Load saved conversations
if "all_conversations" not in st.session_state:
    # Use cache if available (preloaded from Home page)
    if "module2_conversations_cache" in st.session_state:
        saved_conversations = st.session_state.module2_conversations_cache
    else:
        saved_conversations = SessionManager.load_conversations(
            username, st.session_state.guided_session_id
        )
    st.session_state.all_conversations = {
        int(k): v for k, v in saved_conversations.items()
    }

    # Initialize open chat conversation if not present
    if "open_chat" not in st.session_state.all_conversations:
        st.session_state.all_conversations["open_chat"] = []
    
    # Ensure all observations have a conversation list (even if empty)
    for i in range(len(load_question_bank())):
        if i not in st.session_state.all_conversations:
            st.session_state.all_conversations[i] = []

if "guided_phase" not in st.session_state:
    st.session_state.guided_phase = "intro"

if "current_question_idx" not in st.session_state:
    st.session_state.current_question_idx = 0

if "question_bank" not in st.session_state:
    # Use cache if available (preloaded from Home page)
    if "module2_questions_cache" in st.session_state:
        st.session_state.question_bank = st.session_state.module2_questions_cache
    else:
        st.session_state.question_bank = load_question_bank()

if "message_buffer" not in st.session_state:
    st.session_state.message_buffer = MessageBuffer()

if "open_chat_mode" not in st.session_state:
    st.session_state.open_chat_mode = False

if "timer_started" not in st.session_state:
    st.session_state.timer_started = False

# Load RAG documents
rag_status = load_rag_documents()
if "rag_load_status" not in st.session_state:
    st.session_state.rag_load_status = rag_status
    analytics.rag_load_status(
        username,
        st.session_state.guided_session_id,
        rag_status["reference"],
        rag_status["supporting"],
        rag_status["total"],
        rag_status["errors"],
    )

# ==================== STATE PERSISTENCE ====================
def sync_session_to_storage():
    """Sync current session state to persistent storage."""
    try:
        # Save conversations
        SessionManager.save_conversations(
            username,
            st.session_state.guided_session_id,
            st.session_state.all_conversations,
        )

        # Update metadata
        metadata = SessionManager.load_session_metadata(
            username, st.session_state.guided_session_id
        )
        if metadata:
            metadata.update(
                {
                    "phase": st.session_state.guided_phase,
                    "current_question_idx": st.session_state.current_question_idx,
                    "total_messages": sum(
                        len(conv)
                        for conv in st.session_state.all_conversations.values()
                    ),
                    "last_activity": datetime.now().isoformat(),
                }
            )
            SessionManager.save_session_metadata(
                username, st.session_state.guided_session_id, metadata
            )
    except Exception as e:
        logger.error(f"Error syncing session: {e}")
        # Calculate elapsed time safely
        if st.session_state.guided_session_start is not None:
            elapsed = (datetime.now() - st.session_state.guided_session_start).total_seconds()
        else:
            elapsed = 0
        analytics.error_occurred(
            username,
            st.session_state.guided_session_id,
            type(e).__name__,
            str(e),
            elapsed,
            {"context": "sync_session_to_storage"},
        )


# ==================== NAVIGATION HANDLER ====================
def handle_navigation(
    target_question_idx: int, target_phase: str, log_event: bool = True
) -> Tuple[bool, Optional[str]]:
    """
    Handle navigation with server-side validation.

    Returns:
        (success, error_message)
    """
    current_phase = st.session_state.guided_phase
    current_idx = st.session_state.current_question_idx

    # Validate
    is_valid, error_msg = NavigationValidator.validate_navigation(
        current_phase,
        current_idx,
        target_phase,
        target_question_idx,
        len(st.session_state.question_bank),
    )

    if not is_valid:
        return False, error_msg

    # Record timing with defensive check
    if st.session_state.guided_session_start is not None:
        elapsed = (datetime.now() - st.session_state.guided_session_start).total_seconds()
    else:
        elapsed = 0

    # Log event if moving observations
    if target_question_idx != current_idx and log_event:
        analytics.observation_advanced(
            username,
            st.session_state.guided_session_id,
            current_idx,
            target_question_idx,
            elapsed,
            len(st.session_state.all_conversations[current_idx]),
            elapsed,  # TODO: track per-observation timing
        )

    # Apply navigation
    st.session_state.current_question_idx = target_question_idx
    if target_phase != current_phase:
        analytics.phase_transition(
            username,
            st.session_state.guided_session_id,
            current_phase,
            target_phase,
            elapsed,
            {"to_question": target_question_idx},
        )
        st.session_state.guided_phase = target_phase

    sync_session_to_storage()
    return True, None


# ==================== CALCULATE TIME ====================

# Timer only starts after user says "start"
if st.session_state.guided_session_start is None:
    elapsed = timedelta(seconds=0)
    remaining = SESSION_DURATION
    time_expired = False
else:
    elapsed = datetime.now() - st.session_state.guided_session_start
    remaining = SESSION_DURATION - elapsed
    time_expired = remaining.total_seconds() <= 0

# Initialize warning states
if "time_warning_2min_shown" not in st.session_state:
    st.session_state.time_warning_2min_shown = False

# Auto-transition to expired if time is up and in active phase
if time_expired and st.session_state.guided_phase == "active":
    st.session_state.guided_phase = "expired"
    elapsed_seconds = elapsed.total_seconds()
    
    analytics.session_time_expired(
        username,
        st.session_state.guided_session_id,
        st.session_state.current_question_idx,
        st.session_state.current_question_idx,
        sum(len(c) for c in st.session_state.all_conversations.values()),
        elapsed_seconds,
    )
    sync_session_to_storage()

# 2-minute warning
if (remaining.total_seconds() <= 120 and 
    remaining.total_seconds() > 0 and 
    not st.session_state.time_warning_2min_shown and
    st.session_state.guided_phase in ("active", "review")):
    
    st.session_state.time_warning_2min_shown = True
    st.warning(
        "**2 MINUTES REMAINING!**\n\n"
        "Your session will end in 2 minutes. "
        "Consider wrapping up your current thoughts."
    )
    analytics.log_event("two_minute_warning", {
        "remaining_seconds": remaining.total_seconds(),
        "observation_idx": st.session_state.current_question_idx
    }) if hasattr(analytics, 'log_event') else None

# ==================== SIDEBAR WITH TIMER ====================
with st.sidebar:
    st.markdown(f"**Username:** {username}")
    
    # ==================== RENDER TIMER IN SIDEBAR ====================
    countdown_timer(should_start=st.session_state.timer_started)

    # Completion checkbox
    if "completion_status" not in st.session_state:
        st.session_state["completion_status"] = {}

    persistent_value = st.session_state["completion_status"].get(
        "guided_interaction", False
    )
    st.session_state["include_guided_interaction"] = persistent_value

    def _on_include_guided_change():
        current_value = st.session_state.get("include_guided_interaction", False)
        st.session_state["completion_status"]["guided_interaction"] = current_value

    st.checkbox(
        "Check this when done",
        key="include_guided_interaction",
        on_change=_on_include_guided_change,
    )
    
    # Reminder when time is running low or expired
    if time_expired or (remaining.total_seconds() < 300 and st.session_state.guided_phase in ("active", "review")):
        st.info(
            "**Reminder:** Don't forget to check the "
            "**'Check this when done'** checkbox above when you finish!"
        )
    
    # Start button (only in intro phase)
    if st.session_state.guided_phase == "intro":
        if st.button("▶️ Start", use_container_width=True, key="start_button", type="primary"):
            st.session_state.guided_session_start = datetime.now()
            st.session_state.guided_phase = "active"
            st.session_state.current_question_idx = 0
            st.session_state.timer_started = True  # Set flag to track timer is started
            elapsed_seconds = 0
            analytics.phase_transition(
                username,
                st.session_state.guided_session_id,
                "intro",
                "active",
                elapsed_seconds,
                {"observations_completed": 0},
            )
            sync_session_to_storage()
            st.rerun()
    
    # Next button (only in active phase) / End and Save button (only in review phase)
    if st.session_state.guided_phase == "active":
        if st.button("⏭️ Next", use_container_width=True, key="next_button", type="primary"):
            current_idx = st.session_state.current_question_idx
            next_idx = current_idx + 1
            
            if next_idx >= len(st.session_state.question_bank):
                success, error = handle_navigation(current_idx, "review", log_event=False)
            else:
                success, error = handle_navigation(next_idx, "active")
            
            if not success:
                st.error(f"Cannot proceed: {error}")
            else:
                st.rerun()
    
    elif st.session_state.guided_phase == "review":
        if st.button("✓ End Session & Save Data", use_container_width=True, key="end_session_button", type="primary"):
            st.session_state.guided_phase = "complete"
            sync_session_to_storage()
            st.rerun()

    # Show reference conversation
    with st.expander("Show Referenced Full Conversation", expanded=True):
        ref_conversation = load_reference_conversation()
        if ref_conversation:
            for i, turn in enumerate(ref_conversation, 1):
                line_num = i  # 1-indexed line numbers
                is_client = turn.strip().startswith("Client: ")
                # Extract speaker and content
                speaker = "Client:" if is_client else "Therapist:"
                content = turn.replace(f"{speaker} ", "", 1).strip()
                
                if is_client:
                    # Right-justify client's messages with custom CSS
                    st.markdown(f"""
                    <div style="text-align: right; margin-left: 0%; padding: 10px; border-radius: 10px;">
                    <small style="color: #888; font-size: 0.8em;">[Line {line_num}] - {speaker}</small><br>
                    {content}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    # Italicize therapist's messages, but keep line numbers unbolded
                    st.markdown(f"<div style='font-weight:600; font-size:1.08em; margin: 10px 0;'><small style='color: #888; font-size: 0.8em; font-weight: normal;'>[Line {line_num}] - {speaker}</small><br><em>{content}</em></div>", unsafe_allow_html=True)
        else:
            st.warning("Reference conversation not loaded")


# ==================== MAIN CONTENT ====================

st.title("Module 2")
st.markdown(
    "<p style='font-size:12px;color:#e11d48;margin-top:6px;'>"
    "<strong>Privacy Reminder:</strong> Please do not include any identifying information in your messages.</p>",
    unsafe_allow_html=True,
)
st.markdown(
    "<p style='font-size:12px;color:#6b7280;margin-bottom:6px;'>"
    "Disclaimer: TeamMait may be incorrect or incomplete. Please verify information.</p>",
    unsafe_allow_html=True,
)

# Show RAG load warnings
if (
    not st.session_state.rag_load_status["reference"]
    or st.session_state.rag_load_status["total"] < 5
):
    st.warning(
        "**Limited context available**: "
        "Reference transcript may not be fully loaded. AI responses may be less informed."
    )

# ==================== INTRO PHASE ====================
def get_intro_message():
    return """
### Help:
In this Module, I'll share **3 structured observations** about the therapy session.

### How to use:
1. **Read each observation**
2. **Discuss, skip, and/or advance** - You can:
   - Take time to read the transcript, if you wish.
   - Type a question or comment to discuss the observation further.
   - Click the **⏭️ Next** button to move to the next observation.
   - Feel free to skip observations if it does not interest you.
3. **Review phase**
   - Once you've reviewed all 3 observations, you'll enter the review phase.
   - You can continue discussing with TeamMait about any aspect of the session.
   - Click **✓ End Session & Save Data** when you're ready to conclude.
4. **Time limit** - You have <u>**20 minutes total**</u> for the entire session.

### Note:
- You can only move **forward** through observations.
"""

if st.session_state.guided_phase == "intro":
    # Store that intro has been shown to preserve content
    if "intro_shown" not in st.session_state:
        st.session_state.intro_shown = True
    
    st.markdown( get_intro_message(), unsafe_allow_html=True )
    st.info(
            "**Ready to begin?** Click the **▶️ Start** button in the side panel."
        )
    
elif st.session_state.guided_phase == "active":
    # reprint the intro
    st.markdown( get_intro_message(), unsafe_allow_html=True )

    # Show all COMPLETED observations first (with their conversations)
    for obs_idx in range(st.session_state.current_question_idx):
        obs = st.session_state.question_bank[obs_idx]
        
        st.divider()

        st.markdown(f"### Item {obs_idx + 1} of 3 (Completed):")
        
        with st.container(border=True):
            render_feedback_item(obs)
                
        # Show all messages for this completed observation
        for msg in st.session_state.all_conversations[obs_idx]:
            timestamp = msg.get("timestamp", "")
            time_str = ""
            if timestamp:
                try:
                    dt = datetime.fromisoformat(timestamp)
                    time_str = dt.strftime("%H:%M")
                except:
                    pass
            
            with st.chat_message(msg["role"]):
                if time_str:
                    st.caption(f"_{time_str}_")
                st.markdown(msg["content"])
        
    
    # Now show the CURRENT observation
    if st.session_state.current_question_idx < len(st.session_state.question_bank):
        current_q = st.session_state.question_bank[
            st.session_state.current_question_idx
        ]
        current_idx = st.session_state.current_question_idx

        st.divider()

        st.markdown(f"### Item {current_idx + 1} of 3:")

        with st.container(border=True):
            render_feedback_item(current_q)

        st.info(
            "Type a response to discuss this item, "
            "or use the **⏭️ Next** button to move to the next one."
        )

        # Display conversation history for current observation
        for msg in st.session_state.all_conversations[current_idx]:
            timestamp = msg.get("timestamp", "")
            time_str = ""
            if timestamp:
                try:
                    dt = datetime.fromisoformat(timestamp)
                    time_str = dt.strftime("%H:%M")
                except:
                    pass
            
            with st.chat_message(msg["role"]):
                if time_str:
                    st.caption(f"_{time_str}_")
                st.markdown(msg["content"])
        
        # Auto-scroll to latest message
        st.markdown("""
        <script>
            window.scrollTo(0, document.body.scrollHeight);
        </script>
        """, unsafe_allow_html=True)
        
        # User input
        user_input = st.chat_input("Your response or question...")
        
        if user_input:
            # Check if time expired while user was typing
            if time_expired:
                st.error("Session time has expired. Your response could not be saved.")
                st.session_state.guided_phase = "expired"
                sync_session_to_storage()
                # st.rerun()

            # Check for navigation intent first
            if InputParser.detect_navigation_intent(user_input):
                st.info(InputParser.get_navigation_redirect_message())
            else:
                # Log user message - handle None start time
                if st.session_state.guided_session_start is None:
                    elapsed_seconds = 0
                else:
                    elapsed_seconds = (
                        datetime.now() - st.session_state.guided_session_start
                    ).total_seconds()
                
                analytics.user_message(
                    username,
                    st.session_state.guided_session_id,
                    current_idx,
                    elapsed_seconds,
                    len(user_input),
                    "message",
                )

                # Check for duplicates and near-duplicates
                is_new, is_near_duplicate = st.session_state.message_buffer.add_message(user_input)
                
                if not is_new:
                    st.warning("That looks like the same message. Please type something new.")
                elif is_near_duplicate:
                    st.info("ℹ️ You asked something very similar to your last question. I already provided an answer above. Would you like me to expand on it, or do you have a different question?")
                else:
                    # Add to history WITH TIMESTAMP
                    st.session_state.all_conversations[current_idx].append(
                        {
                            "role": "user",
                            "content": user_input,
                            "timestamp": datetime.now().isoformat()
                        }
                    )

                    # Immediately add a placeholder assistant message showing "Thinking..."
                    st.session_state.all_conversations[current_idx].append(
                        {
                            "role": "assistant",
                            "content": "*Thinking...*",
                            "timestamp": datetime.now().isoformat()
                        }
                    )
                    
                    sync_session_to_storage()
                    st.rerun()  # Rerun to display both messages immediately


# Continue processing after rerun if we have a thinking message
current_idx = st.session_state.current_question_idx
if st.session_state.guided_phase == "active" and current_idx < len(st.session_state.question_bank) and st.session_state.all_conversations[current_idx]:
    if st.session_state.all_conversations[current_idx][-1]["content"] == "*Thinking...*":
        # Get the user message (it's the one before thinking)
        user_input = st.session_state.all_conversations[current_idx][-2]["content"] if len(st.session_state.all_conversations[current_idx]) > 1 else ""
        
        # Generate AI response to replace the thinking message
        current_q_data = st.session_state.question_bank[current_idx]
        context = retrieve_context(user_input)
        observation_context = format_observation_context(current_q_data)
        system_prompt = (
            build_system_prompt()
            + f"\n\nCurrent Observation Being Discussed:\n{observation_context}\n\n"
            f"Context from transcript:\n{context}"
        )

        try:
            # Generate response
            acc = ""
            start_time = time.time()

            response_gen = OpenAIHandler.openai_complete(
                history=st.session_state.all_conversations[current_idx],
                system_text=system_prompt,
                client=client,
                stream=True,
                max_tokens=512,
                max_retries=2,
                timeout=30,
            )

            for chunk in response_gen:
                acc += chunk

            generation_time = time.time() - start_time

            # Replace the thinking message with the actual response
            st.session_state.all_conversations[current_idx][-1] = {
                "role": "assistant",
                "content": acc.strip(),
                "timestamp": datetime.now().isoformat()
            }

            # Log response
            if st.session_state.guided_session_start is not None:
                elapsed_seconds = (
                    datetime.now() - st.session_state.guided_session_start
                ).total_seconds()
            else:
                elapsed_seconds = 0
            tokens_estimated = len(acc) // 4
            analytics.ai_response(
                username,
                st.session_state.guided_session_id,
                current_idx,
                elapsed_seconds,
                len(acc),
                tokens_estimated,
                generation_time,
            )

            sync_session_to_storage()
            # Don't rerun - let message display on next timer tick

        except (APIRetryableError, APIPermanentError) as e:
            error_msg = OpenAIHandler.format_error_message(e)
            st.error(error_msg)

            # Remove both user and thinking messages on error
            st.session_state.all_conversations[current_idx].pop()
            st.session_state.all_conversations[current_idx].pop()
            sync_session_to_storage()

            if st.session_state.guided_session_start is not None:
                elapsed_seconds = (
                    datetime.now() - st.session_state.guided_session_start
                ).total_seconds()
            else:
                elapsed_seconds = 0
            analytics.error_occurred(
                username,
                st.session_state.guided_session_id,
                type(e).__name__,
                str(e),
                elapsed_seconds,
                {"context": "ai_response", "observation_idx": current_idx},
            )

            logger.error(f"API error in observation {current_idx}: {e}")

        except Exception as e:
            error_msg = "Unexpected error. Please try again."
            st.error(error_msg)

            # Remove both user and thinking messages on error
            st.session_state.all_conversations[current_idx].pop()
            st.session_state.all_conversations[current_idx].pop()
            sync_session_to_storage()

            if st.session_state.guided_session_start is not None:
                elapsed_seconds = (
                    datetime.now() - st.session_state.guided_session_start
                ).total_seconds()
            else:
                elapsed_seconds = 0
            analytics.error_occurred(
                username,
                st.session_state.guided_session_id,
                type(e).__name__,
                str(e),
                elapsed_seconds,
                {"context": "ai_response", "observation_idx": current_idx},
            )

            logger.error(
                f"Unexpected error in observation {current_idx}: {e}",
                exc_info=True,
            )

    # Don't auto-transition to review - user must click Next button

# ==================== TIME EXPIRED PHASE ====================

elif st.session_state.guided_phase == "expired":
    # Initialize warning states
    if "time_warning_2min_shown" not in st.session_state:
        st.session_state.time_warning_2min_shown = False
    
    current_idx = min(
        st.session_state.current_question_idx,
        len(st.session_state.question_bank) - 1,
    )
    
    st.error("### Session Time Expired")
    
    st.markdown(
        """
    Your 20-minute session has ended.
    
    ### Your Discussion is Saved
    
    All your messages below are preserved and will be included in your results.
    
    ### Final Thoughts
    
    You can share **final thoughts** below for a brief synthesis from TeamMait.
    
    Or skip directly to the next step.
    """
    )
    
    st.divider()
    
    # Show previous messages with timestamps (NEVER ERASE!)
    st.markdown("### Your Previous Discussion")
    
    if st.session_state.all_conversations[current_idx]:
        for msg in st.session_state.all_conversations[current_idx]:
            timestamp = msg.get("timestamp", "")
            time_str = ""
            if timestamp:
                try:
                    dt = datetime.fromisoformat(timestamp)
                    time_str = dt.strftime("%H:%M")
                except:
                    pass
            
            with st.chat_message(msg["role"]):
                if time_str:
                    st.caption(f"_{time_str}_")
                st.markdown(msg["content"])
    else:
        st.info("No messages yet")
    
    st.divider()
    
    # Input for final thoughts
    st.markdown("### Final Thoughts (Optional)")
    user_input = st.chat_input(
        "Type your final thoughts...",
        placeholder="Leave blank to skip to next step"
    )
    
    if user_input:
        # Add user message with timestamp
        st.session_state.all_conversations[current_idx].append(
            {
                "role": "user",
                "content": user_input,
                "timestamp": datetime.now().isoformat()
            }
        )
        
        with st.chat_message("user"):
            st.caption(f"_{datetime.now().strftime('%H:%M')}_")
            st.markdown(user_input)
        
        # Generate synthesis (NOT summary)
        with st.chat_message("assistant"):
            placeholder = st.empty()
            placeholder.markdown("*Synthesizing your insights...*")
            
            try:
                # Include the observation context for synthesis
                current_q_data = st.session_state.question_bank[current_idx]
                observation_context = format_observation_context(current_q_data)
                system_prompt = (
                    build_system_prompt()
                    + f"\n\nCurrent Observation:\n{observation_context}\n\n"
                    "Provide a brief synthesis of the key insights and observations "
                    "from this discussion. Do NOT summarize what was said. Instead, "
                    "focus on: What was learned? What patterns emerged? What stands out? "
                    "Keep it to 2-3 sentences maximum. Be insightful, not repetitive."
                )
                
                acc = ""
                start_time = time.time()
                first_chunk = True
                
                response_gen = OpenAIHandler.openai_complete(
                    history=st.session_state.all_conversations[current_idx],
                    system_text=system_prompt,
                    client=client,
                    stream=True,
                    max_tokens=256,
                    max_retries=2,
                )
                
                for chunk in response_gen:
                    if first_chunk:
                        placeholder.empty()
                        first_chunk = False
                    
                    acc += chunk
                    placeholder.markdown(acc.lstrip())
                
                # Save synthesis with timestamp
                st.session_state.all_conversations[current_idx].append(
                    {
                        "role": "assistant",
                        "content": acc.strip(),
                        "timestamp": datetime.now().isoformat()
                    }
                )
                
                sync_session_to_storage()
                
            except Exception as e:
                placeholder.error(
                    f"Could not generate synthesis: {str(e)}"
                )
                logger.error(f"Synthesis generation error: {e}")
    
    # Buttons
    st.divider()
    st.markdown("### Next Step")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button(
            "Skip Final Thoughts",
            use_container_width=True,
            key="skip_final"
        ):
            st.session_state.guided_phase = "complete"
            sync_session_to_storage()
            st.rerun()
    
    with col2:
        if st.button(
            "Continue to Results",
            type="primary",
            use_container_width=True,
            key="continue_after_expired"
        ):
            st.session_state.guided_phase = "complete"
            sync_session_to_storage()
            st.rerun()
    
    # Reminder
    st.info(
        "**Reminder:** Don't forget to check the "
        "**'Check this when done'** checkbox in the sidebar!"
    )

# ==================== REVIEW PHASE ====================
elif st.session_state.guided_phase == "review":
    # Clear the open_chat conversation on first entry to review phase
    if "review_phase_entered" not in st.session_state:
        st.session_state.all_conversations["open_chat"] = []
        st.session_state.review_phase_entered = True
        sync_session_to_storage()
    
    # Recalculate remaining time using the original session start
    adjusted_remaining = COUNTDOWN_DURATION_SECONDS - (datetime.now() - st.session_state.guided_session_start).total_seconds()
    
    # Automatically skip to open chat mode if there's time, or finish if time expired
    if time_expired or adjusted_remaining <= 0:
        # Time expired - go straight to complete phase
        st.error("### Session Time Expired")
        st.markdown("Thank you for completing all observations. The module is now complete.")
        
        if st.button("Finish Module", type="primary", use_container_width=True):
            st.session_state.guided_phase = "complete"
            sync_session_to_storage()
            st.rerun()
    else:
        # Display summary of all prior observations and their chat histories
        st.markdown("## Review: Discussion Summary")
        st.divider()
        
        for obs_idx in range(len(st.session_state.question_bank)):
            obs_item = st.session_state.question_bank[obs_idx]
            obs_title = obs_item.get("title", f"Observation {obs_idx + 1}")
            
            with st.expander(f"**Observation {obs_idx + 1}**: {obs_title}", expanded=False):
                # First, display the original observation
                st.markdown("**Original Observation:**")
                render_feedback_item(obs_item)
                
                st.divider()
                
                st.markdown("**Discussion History:**")
                # Display conversation history for this observation
                if st.session_state.all_conversations[obs_idx]:
                    for msg in st.session_state.all_conversations[obs_idx]:
                        timestamp = msg.get("timestamp", "")
                        time_str = ""
                        if timestamp:
                            try:
                                dt = datetime.fromisoformat(timestamp)
                                time_str = dt.strftime("%H:%M")
                            except:
                                pass
                        
                        with st.chat_message(msg["role"]):
                            if time_str:
                                st.caption(f"_{time_str}_")
                            st.markdown(msg["content"])
                else:
                    st.info("No discussion for this observation.")
        
        st.divider()
        
        # Generate key points summary on first arrival if not already done
        if "observations_summary_generated" not in st.session_state:
            summary = generate_observations_summary(st.session_state.all_conversations, client)
            if summary:
                st.session_state.observations_summary = summary
                st.session_state.observations_summary_generated = True
                sync_session_to_storage()
        
        # Open chat mode - always shown in review phase with time remaining
        current_idx = "open_chat"
        
        st.markdown("### Continue the Conversation")
        
        # Display the summary as markdown (not as a chat message)
        if st.session_state.observations_summary_generated and st.session_state.observations_summary:
            st.markdown("**Key Takeaways from Your Discussion:**")
            st.markdown(st.session_state.observations_summary)
        
        st.divider()
        
        # Prompt for continued discussion
        remaining_time_sec = max(0, int(adjusted_remaining))
        remaining_min = remaining_time_sec // 60
        remaining_sec = remaining_time_sec % 60
        
        st.info(
            f"You have **{remaining_min}:{remaining_sec:02d}** remaining. "
            "Feel free to discuss any other aspects of the session, ask questions, or share additional thoughts. "
            "When you're done, check the **'Check this when done'** box to save your results."
        )
        
        # Display conversation history WITH TIMESTAMPS
        for msg in st.session_state.all_conversations[current_idx]:
            timestamp = msg.get("timestamp", "")
            time_str = ""
            if timestamp:
                try:
                    dt = datetime.fromisoformat(timestamp)
                    time_str = dt.strftime("%H:%M")
                except:
                    pass
            
            with st.chat_message(msg["role"]):
                if time_str:
                    st.caption(f"_{time_str}_")
                st.markdown(msg["content"])
        
        # Auto-scroll to latest message
        st.markdown("""
        <script>
            window.scrollTo(0, document.body.scrollHeight);
        </script>
        """, unsafe_allow_html=True)
        
        # User input
        user_input = st.chat_input("Your response or question...")
        
        if user_input:
            # Check if time expired while user was typing
            if adjusted_remaining <= 0:
                st.error("Session time has expired. Your response could not be saved.")
                st.session_state.guided_phase = "expired"
                sync_session_to_storage()
            else:
                # Log user message - handle None start time
                if st.session_state.guided_session_start is None:
                    elapsed_seconds = 0
                else:
                    elapsed_seconds = (
                        datetime.now() - st.session_state.guided_session_start
                    ).total_seconds()
                
                analytics.user_message(
                    username,
                    st.session_state.guided_session_id,
                    current_idx,
                    elapsed_seconds,
                    len(user_input),
                    "message",
                )
                
                # Check for duplicates and near-duplicates
                is_new, is_near_duplicate = st.session_state.message_buffer.add_message(user_input)
                
                if not is_new:
                    st.warning("That looks like the same message. Please type something new.")
                elif is_near_duplicate:
                    st.info("ℹ️ You asked something very similar to your last question. I already provided an answer above. Would you like me to expand on it, or do you have a different question?")
                else:
                    # Add to history WITH TIMESTAMP
                    st.session_state.all_conversations[current_idx].append(
                        {
                            "role": "user",
                            "content": user_input,
                            "timestamp": datetime.now().isoformat()
                        }
                    )
                    
                    with st.chat_message("user"):
                        st.caption(f"_{datetime.now().strftime('%H:%M')}_")
                        st.markdown(user_input)
                    
                    # Generate AI response
                    context = retrieve_context(user_input)
                    # Include all observation context in review phase
                    observation_contexts = []
                    for i, q in enumerate(st.session_state.question_bank):
                        obs_formatted = format_observation_context(q)
                        observation_contexts.append(f"Observation {i+1}: {obs_formatted}")
                    all_observations = "\n\n".join(observation_contexts)
                    system_prompt = (
                        build_system_prompt()
                        + f"\n\nAll Observations from This Supervision Session:\n{all_observations}\n\n"
                        f"Context from transcript:\n{context}"
                    )
                    
                    with st.chat_message("assistant"):
                        placeholder = st.empty()
                        placeholder.markdown("*Thinking...*")
                        
                        try:
                            # Generate response
                            acc = ""
                            start_time = time.time()
                            first_chunk = True
                            
                            response_gen = OpenAIHandler.openai_complete(
                                history=st.session_state.all_conversations[current_idx],
                                system_text=system_prompt,
                                client=client,
                                stream=True,
                                max_tokens=512,
                                max_retries=2,
                                timeout=30,
                            )
                            
                            for chunk in response_gen:
                                if first_chunk:
                                    placeholder.empty()
                                    first_chunk = False
                                
                                acc += chunk
                                placeholder.markdown(acc.lstrip())
                            
                            generation_time = time.time() - start_time
                            
                            # Save to history WITH TIMESTAMP
                            st.session_state.all_conversations[current_idx].append(
                                {
                                    "role": "assistant",
                                    "content": acc.strip(),
                                    "timestamp": datetime.now().isoformat()
                                }
                            )
                            
                            # Log response
                            if st.session_state.guided_session_start is not None:
                                elapsed_seconds = (
                                    datetime.now() - st.session_state.guided_session_start
                                ).total_seconds()
                            else:
                                elapsed_seconds = 0
                            tokens_estimated = len(acc) // 4
                            analytics.ai_response(
                                username,
                                st.session_state.guided_session_id,
                                current_idx,
                                elapsed_seconds,
                                len(acc),
                                tokens_estimated,
                                generation_time,
                            )
                            
                            sync_session_to_storage()
                            st.rerun()
                        
                        except (APIRetryableError, APIPermanentError) as e:
                            error_msg = OpenAIHandler.format_error_message(e)
                            placeholder.error(error_msg)
                            
                            # Remove the user message to keep state clean
                            st.session_state.all_conversations[current_idx].pop()
                            sync_session_to_storage()
                            
                            if st.session_state.guided_session_start is not None:
                                elapsed_seconds = (
                                    datetime.now() - st.session_state.guided_session_start
                                ).total_seconds()
                            else:
                                elapsed_seconds = 0
                            analytics.error_occurred(
                                username,
                                st.session_state.guided_session_id,
                                type(e).__name__,
                                str(e),
                                elapsed_seconds,
                                {"context": "ai_response", "observation_idx": current_idx},
                            )
                            
                            logger.error(f"API error in open chat: {e}")
                        
                        except Exception as e:
                            error_msg = "Unexpected error. Please try again."
                            placeholder.error(error_msg)
                            
                            # Remove the user message
                            st.session_state.all_conversations[current_idx].pop()
                            sync_session_to_storage()
                            
                            if st.session_state.guided_session_start is not None:
                                elapsed_seconds = (
                                    datetime.now() - st.session_state.guided_session_start
                                ).total_seconds()
                            else:
                                elapsed_seconds = 0
                            analytics.error_occurred(
                                username,
                                st.session_state.guided_session_id,
                                type(e).__name__,
                                str(e),
                                elapsed_seconds,
                                {"context": "ai_response", "observation_idx": current_idx},
                            )
                            
                            logger.error(
                                f"Unexpected error in open chat: {e}",
                                exc_info=True,
                            )
        
        # Use Streamlit's sleep to auto-rerun and keep timer updating (but don't rerun in review immediately)
        import time as time_module
        time_module.sleep(1)
        st.rerun()

# ==================== COMPLETE PHASE ====================

elif st.session_state.guided_phase == "complete":
    reviewed = min(st.session_state.current_question_idx, 4)
    total_messages = sum(len(c) for c in st.session_state.all_conversations.values())
    # Defensive check: if session_start is None, default to 0 elapsed time
    if st.session_state.guided_session_start is not None:
        elapsed_seconds = (datetime.now() - st.session_state.guided_session_start).total_seconds()
    else:
        elapsed_seconds = 0

    st.success("### Module Complete")
    st.markdown(
        f"""
    Thank you for completing this module!

    **Summary:**
    - **Items reviewed:** {reviewed} / 3
    - **Total messages:** {total_messages}
    - **Time used:** {int(elapsed_seconds // 60)}:{int(elapsed_seconds % 60):02d}
    """
    )

    # Log completion
    analytics.session_completed(
        username,
        st.session_state.guided_session_id,
        reviewed,
        total_messages,
        elapsed_seconds,
        status="success",
    )

    # Mark session as complete in storage
    SessionManager.complete_session(
        username,
        st.session_state.guided_session_id,
        {
            "observations_reviewed": reviewed,
            "total_messages": total_messages,
            "elapsed_seconds": elapsed_seconds,
        },
    )

    # Clean up message buffer
    st.session_state.message_buffer.clear()

    if st.button("Continue to Next Step", type="primary", use_container_width=True):
        st.switch_page("pages/4_Finish.py")


# ==================== AUTO-RERUN TIMER ====================
# Rerun every second if timer is active to keep it updating
if st.session_state.timer_started and st.session_state.guided_session_start is not None and st.session_state.guided_phase == "active":
    elapsed = (datetime.now() - st.session_state.guided_session_start).total_seconds()
    remaining_time = COUNTDOWN_DURATION_SECONDS - elapsed
    if 0 < remaining_time <= COUNTDOWN_DURATION_SECONDS:
        # Use st.empty() to trigger rerun without blocking
        import time as time_module
        time_module.sleep(0.5)
        st.rerun()