import streamlit as st
import time
from textwrap import dedent
from urllib.parse import quote
import json
from datetime import datetime
import os
from oauth2client.service_account import ServiceAccountCredentials
import gspread
from sentence_transformers import SentenceTransformer
import sys
import glob
from utils.input_parser import MessageBuffer

# ---------- SQLite shim for Chroma ----------
try:
    import pysqlite3
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    pass

import chromadb
from chromadb.utils import embedding_functions
from chromadb.config import Settings, DEFAULT_TENANT, DEFAULT_DATABASE

# Optional SDKs (load if installed)
try:
    import anthropic
except ImportError:
    anthropic = None
try:
    from openai import OpenAI  # noqa: F401
except ImportError:
    OpenAI = None
try:
    from docx import Document
except ImportError:
    Document = None

# ---------- Page & Theme ----------
st.set_page_config(page_title="Open Chat", page_icon="", layout="wide")

# ---------- Simple avatars (SVG data URIs) ----------
from urllib.parse import quote as _quote  # noqa: E402
def svg_data_uri(svg: str) -> str:
    return "data:image/svg+xml;utf8," + _quote(svg)

DM_SVG = svg_data_uri(
    dedent("""
    <svg xmlns='http://www.w3.org/2000/svg' width='64' height='64'>
      <defs>
        <linearGradient id='g' x1='0' x2='1' y1='0' y2='1'>
          <stop offset='0%' stop-color='#1f2937'/>
          <stop offset='100%' stop-color='#0b1220'/>
        </linearGradient>
      </defs>
      <circle cx='32' cy='32' r='32' fill='url(#g)'/>
      <text x='50%' y='54%' text-anchor='middle' font-family='Inter,Arial' font-size='24' fill='#e5e7eb' font-weight='700'>U</text>
    </svg>
    """).strip()
)

BOT_SVG = svg_data_uri(
    dedent("""
    <svg xmlns='http://www.w3.org/2000/svg' width='64' height='64'>
      <circle cx='32' cy='32' r='32' fill='#111827'/>
      <rect x='18' y='20' width='28' height='22' rx='6' fill='#1f2937' stroke='#374151' stroke-width='2'/>
      <circle cx='26' cy='31' r='4' fill='#93c5fd'/>
      <circle cx='38' cy='31' r='4' fill='#93c5fd'/>
      <rect x='28' y='42' width='8' height='6' rx='3' fill='#4b5563'/>
      <rect x='30' y='12' width='4' height='8' rx='2' fill='#6b7280'/>
    </svg>
    """).strip()
)

# ---------- Helpers ----------
def now_ts() -> str:
    return datetime.now().isoformat()  # Full ISO format for storage

# ---------- Login dialog ----------
@st.dialog("Login", dismissible=False, width="small")
def get_user_details():
    username = st.text_input("Username")
    password = st.text_input("password")
    submit = st.button("Submit", type="primary")
    if submit:
        if not username or not password:
            st.warning("Please enter both a username and a password.")
            return
        st.session_state.user_info = {"username": username, "password": password}
        # also mirror to top-level keys for convenience
        st.session_state["username"] = username
        st.session_state["password"] = password
        st.rerun()

if "user_info" not in st.session_state:
    get_user_details()
    st.stop()  # <- block the rest of the script until we have creds

# ---------- Now that we have user_info, continue ----------
username = st.session_state["user_info"]["username"]
password = st.session_state["user_info"]["password"]

# ---------- Databasing (quick and sloppy) ----------
scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
creds_dict = json.loads(st.secrets["GOOGLE_CREDENTIALS"])
creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
gs_client = gspread.authorize(creds)
sheet = gs_client.open(st.secrets["SHEET_NAME"]).sheet1

# ---------- Session state ----------
if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "assistant",
            "content": "Hi, my name is TeamMait. Feel free to ask me any questions to ask me about the referenced session transcript, which may be found in the left side panel.",
            "ts": now_ts(),
            "display_name": "TeamMait",
        }
    ]

if "message_buffer" not in st.session_state:
    st.session_state.message_buffer = MessageBuffer()
if "errors" not in st.session_state:
    st.session_state.errors = []

# # ---- Brevity policy & templates (1–5) ----
# BREVITY = {
#     1: dict(name="Detailed Narrative",  max_tokens=900,  bullets=6,  sections=["Strengths","Areas for Growth","Opportunities","Risks/Concerns"], headline=False),
#     2: dict(name="Structured Summary",  max_tokens=650,  bullets=5,  sections=["Strengths","Areas for Growth","Concerns"],                    headline=False),
#     3: dict(name="Balanced Highlights", max_tokens=450,  bullets=4,  sections=None,                                                             headline=False),
#     4: dict(name="Key Points Only",     max_tokens=250,  bullets=3,  sections=None,                                                             headline=False),
#     5: dict(name="Headline Only",       max_tokens=120,  bullets=0,  sections=None,                                                             headline=True),
# }

# ---------- Sidebar (settings) ----------
with st.sidebar:
    st.markdown(f"**Username:** {username}")
    # st.button("Clear chat", type="secondary", on_click=lambda: st.session_state.update(messages=[]))

    # Settings moved to hidden/default values
    stream_on = True
    show_timestamps = True
    model = r"gpt-4o-mini"

    # st.session_state['empathy'] = empathy
    # st.session_state['brevity'] = brevity
    st.session_state['stream_on'] = stream_on
    st.session_state['show_timestamps'] = show_timestamps
    st.session_state['model'] = model


    # Initialize completion status if it doesn't exist
    if "completion_status" not in st.session_state:
        st.session_state["completion_status"] = {}
    
    # Sync the checkbox state with the persistent completion tracker
    persistent_value = st.session_state["completion_status"].get("open_chat", False)
    st.session_state["include_open_chat"] = persistent_value

    # Checkbox removed - completion status tracked automatically

    # st.divider()
    # ---------- Chroma initialization (after login) ----------
    embed_model = "sentence-transformers/all-MiniLM-L6-v2"
    embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(model_name=embed_model)

    chroma_client = chromadb.PersistentClient(
        path="./rag_store",
        settings=Settings(),
        tenant=DEFAULT_TENANT,
        database=DEFAULT_DATABASE,
    )
    collection = chroma_client.get_or_create_collection("therapy", embedding_function=embedding_fn)

    def extract_text_from_docx(docx_path: str) -> str:
        """Extract text content from a DOCX file."""
        if Document is None:
            return f"[Error: python-docx not installed. Cannot read {os.path.basename(docx_path)}]"
        
        try:
            doc = Document(docx_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            return f"[Error reading {os.path.basename(docx_path)}: {str(e)}]"

    # ---------- Reference Conversation (expander) ----------
    @st.cache_resource
    def load_rag_documents():
        doc_folder = "doc/RAG"
        supporting_folder = os.path.join(doc_folder, "supporting_documents")
        documents = []
        ids = []
        # Load main reference conversation
        ref_path = os.path.join(doc_folder, "116_P8_conversation.json")
        if os.path.exists(ref_path):
            with open(ref_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict) and "full_conversation" in data:
                    for i, turn in enumerate(data["full_conversation"]):
                        documents.append(str(turn))
                        ids.append(f"ref_{i}")
                # elif isinstance(data, list): # we don't want the three-turn parts
                #     for i, item in enumerate(data):
                #         documents.append(str(item))
                #         ids.append(f"ref_{i}")
        # Load .txt and .json files from supporting_documents
        for txt_path in glob.glob(os.path.join(supporting_folder, "*.txt")):
            with open(txt_path, "r", encoding="utf-8") as f:
                content = f.read()
                documents.append(content)
                ids.append(f"supp_txt_{os.path.basename(txt_path)}")
        for json_path in glob.glob(os.path.join(supporting_folder, "*.json")):
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    for i, item in enumerate(data):
                        documents.append(str(item))
                        ids.append(f"supp_json_{os.path.basename(json_path)}_{i}")
                elif isinstance(data, dict):
                    for k, v in data.items():
                        documents.append(f"{k}: {v}")
                        ids.append(f"supp_json_{os.path.basename(json_path)}_{k}")
        
        # Load .docx files from supporting_documents
        for docx_path in glob.glob(os.path.join(supporting_folder, "*.docx")):
            content = extract_text_from_docx(docx_path)
            if content and not content.startswith("[Error"):
                # Split large documents into chunks for better retrieval
                words = content.split()
                chunk_size = 500  # words per chunk
                
                if len(words) <= chunk_size:
                    documents.append(content)
                    ids.append(f"supp_docx_{os.path.basename(docx_path)}")
                else:
                    # Split into chunks
                    for i in range(0, len(words), chunk_size):
                        chunk = " ".join(words[i:i + chunk_size])
                        documents.append(chunk)
                        ids.append(f"supp_docx_{os.path.basename(docx_path)}_chunk_{i//chunk_size + 1}")
            else:
                # Add error message as document for debugging
                documents.append(content)
                ids.append(f"supp_docx_error_{os.path.basename(docx_path)}")
        
        # Seed collection if empty
        if collection.count() == 0 and documents:
            collection.add(documents=documents, ids=ids)
        return documents

    rag_documents = load_rag_documents()
    
    # Show only the reference conversation in the expander
    ref_conversation = []
    ref_path = os.path.join("doc/RAG", "116_P8_conversation.json")
    if os.path.exists(ref_path):
        with open(ref_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict) and "full_conversation" in data:
                ref_conversation = data["full_conversation"]
            elif isinstance(data, list):
                ref_conversation = data
    with st.expander("Show Referenced Full Conversation", expanded=True):
        if ref_conversation:
            for i, turn in enumerate(ref_conversation):
                line_num = i + 1  # 1-indexed line numbers
                is_client = turn.strip().startswith("Client: ")
                # Extract speaker and content
                speaker = "Client:" if is_client else "Therapist:"
                content = turn.replace(f"{speaker} ", "", 1).strip()
                
                if is_client:
                    # Right-justify client's messages with custom CSS
                    st.markdown(f"""
                    <div style="text-align: right; margin-left: 0%; padding: 10px; border-radius: 10px;">
                    <small style="color: #888; font-size: 0.8em;">[Line {line_num}] - {speaker}</small><br>
                    {content}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    # Italicize therapist's messages, but keep line numbers unbolded
                    st.markdown(f"<div style='font-weight:600; font-size:1.08em; margin: 10px 0;'><small style='color: #888; font-size: 0.8em; font-weight: normal;'>[Line {line_num}] - {speaker}</small><br><em>{content}</em></div>", unsafe_allow_html=True)
        else:
            st.info("No reference conversation found in 116_P8_conversation.json.")

# ---------- Layout CSS ----------
st.markdown(
    """
    <style>
      .app-container { max-width: 700px; margin: 0 auto; padding-top: 0px; }
      .chat-wrapper { min-height: 0px; padding: 0 4px; margin-top: 0px; }
      .msg.user { display: flex; justify-content: flex-end; }
      .msg.user .stChatMessage {
            flex-direction: row-reverse;
            display: flex;
            align-items: flex-start;
            margin-right: 0;
            margin-left: auto;
      }
      .stChatMessage {padding-top: 0rem; padding-bottom: 0rem;}
      .stChatInputContainer {position: sticky; bottom: 0; background: var(--background-color);}
      button[data-testid="stDownloadButton"] {
        background-color: #4F8A8B;
        color: red;
        border-radius: 6px;
        border: none;
      }
      button[data-testid="stDownloadButton"]:hover {
        background-color: #306B6B;
      }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown("<div class='app-container'>", unsafe_allow_html=True)
st.title("TeamMait Private Conversation")
st.markdown("<p style='font-size:12px;color:#e11d48;margin-top:6px;'><strong>Privacy Reminder:</strong> Please do not include any identifying information in your messages.</p>", unsafe_allow_html=True)
st.markdown("<p style='font-size:12px;color:#6b7280;margin-bottom:6px;'>Disclaimer: TeamMait may be incorrect or incomplete. Please verify information..</p>", unsafe_allow_html=True)


# # ---------- Load JSON Conversation into Vector Store ----------
# @st.cache_resource
# def load_conversation_and_seed():
#     with open("116_P8_conversation.json") as f:
#         data = json.load(f)
#     if collection.count() == 0:
#         for i, turn in enumerate(data.get("full_conversation", [])):
#             collection.add(documents=[turn], ids=[f"conv_{i}"])
#     return data

# data = load_conversation_and_seed()

# # ---------- Reference Conversation (expander) ----------
# conversation = data.get("full_conversation", [])
# with st.expander("Show Referenced Full Conversation", expanded=False):
#     for turn in conversation:
#         st.markdown(turn)

# ---------- Prompt builder ----------
# def structure_prompt(level:int) -> str:
#     cfg = BREVITY[level]
#     if cfg["headline"]:
#         return (
#             "Output ONE sentence: the single most important clinical takeaway. "
#             "No preamble, no bullets, no quotes."
#         )
#     if cfg["sections"]:
#         # Sectioned bullets for Levels 1–2
#         per = max(1, cfg["bullets"] // len(cfg["sections"]) + 1)
#         sec_lines = "\n".join(f"- {s}: ≤{per} bullets" for s in cfg["sections"])
#         return (
#             "Use sectioned bullets. Each bullet: one actionable point anchored to the transcript; "
#             "start with a strong verb; optional short quote in quotes; ≤25 words per bullet.\n"
#             f"Sections:\n{sec_lines}\n"
#             "No filler, no concluding paragraph."
#         )
#     # Flat bullets for Levels 3–4
#     return (
#         f"Return exactly {cfg['bullets']} bullets. Each bullet ≤25 words, starts with a verb, "
#         "optional (timestamp). No intro or outro."
#     )

# def build_system_prompt(empathy_value: int, brevity_level: int) -> str:
def build_system_prompt() -> str:
    """Build system prompt for open chat mode (user-led discussion).
    
    Matches Module 2's comprehensive framework but without observations/items structure.
    User leads the conversation; TeamMait responds.
    """
    return (
        "You are TeamMait, a peer-support assistant for expert clinicians reviewing therapist performance in a transcript. "
        "Your responses must follow two behavioral modes: global rules (always active) and analysis mode (only when the user requests supervisory analysis).\n\n"

        "1. GLOBAL RULES (ALWAYS ACTIVE)\n"
        "- Never fabricate transcript content, facts, or therapist intentions.\n"
        "- Do not infer internal states, emotions, or off-transcript behavior.\n"
        "- Respond concisely, professionally, and only to what the user asked.\n"
        "- Use a natural, peer-like supervisory tone; avoid rigid sections or templates unless the user requests structure.\n"
        "- Treat each user message independently unless the user explicitly references earlier turns.\n"
        "- If the user gives a dismissive acknowledgment (e.g., \"ok\", \"thanks\", \"got it\"), briefly acknowledge and ask whether they want to continue or move on.\n"
        "- If the user expresses confusion (e.g., \"what?\", \"I don't understand\", \"unclear\"), provide a simpler, more direct explanation of your prior point.\n"
        "- Do not offer unsolicited elaboration or additional insights outside analytic tasks.\n"
        "- Do NOT ask if the user wants you to analyze or offer to analyze. Simply respond to their question or request directly based on context. Only ask for clarification if the user's intent is genuinely ambiguous (e.g., unclear phrasing, conflicting requests).\n"
        "- NEVER end messages with offers like 'Would you like me to...', 'Would you like me to analyze...', 'Should I...', 'Do you want me to...', or similar. This is engagement-seeking behavior. Just respond to what was asked and let the user decide their next move.\n"

        "2. ANALYSIS MODE (ONLY WHEN USER REQUESTS SUPERVISORY ANALYSIS)\n"
        "Enter analysis mode only when the user asks you to analyze therapist behavior, evaluate fidelity, generate observations, or provide supervision-like feedback on the session.\n\n"

        "When in analysis mode, follow these rules:\n\n"

        "Evidence Use:\n"
        "- Cite transcript lines in the format [Line X] when providing evidence.\n"
        "- Base all claims on observable behavior only.\n"
        "- Distinguish clearly in your wording between what is directly observed (evidence) and your interpretation of its relevance to PE fidelity.\n\n"

        "Fidelity Alignment:\n"
        "- Use PE fidelity criteria (e.g., orientation to imaginal exposure, SUDS monitoring, hotspot identification, present-tense prompting, reinforcing comments, processing after imaginal, session structure, off-task discussion) as the interpretive framework.\n"
        "- If the transcript does not provide enough information to evaluate a fidelity domain, explicitly state that the evidence is insufficient.\n\n"

        "Trust Calibration:\n"
        "- Use calibrated language such as \"appears\", \"may indicate\", or \"based on [Line X–Y]\".\n"
        "- Mark uncertainty explicitly; do not overstate confidence. Offer to estimate your certainty in a conclusion.\n\n"

        "Autonomy Preservation:\n"
        "- Frame feedback as observations or suggestions, not directives, unless the user explicitly requests strong prescriptive guidance.\n"
        "- Do not reinterpret or expand the user's goals unless they ask you to.\n\n"

        "Contestability:\n"
        "- Present your reasoning in a way that allows the clinician to agree, disagree, or reinterpret your analysis.\n"
        "- Avoid unverifiable or global statements about therapist competence.\n\n"

        "Boundaries:\n"
        "- Do not generalize beyond this specific transcript or session.\n"
        "- Do not simulate or infer missing dialogue or events not shown in the transcript.\n"
        "- Do not evaluate client behavior or provide therapeutic interpretations of the client.\n\n"

        "Format in Analysis Mode:\n"
        "- Unless the user specifies otherwise, provide ideally 3, but no more than 5, bullet points to answer a given query.\n"
        "- One sentence per bullet; limit each bullet to about 10 words or 75 characters at most (unless the user explicitly requests longer responses).\n"
        "- Use sub-bullets if you can't adequately convey a point in fewer than 10 words or 75 characters. Limit sub-bullets to one per parent-bullet.\n"
        "- Prioritize clarity and brevity; avoid redundancy.\n"
        "- On a scale of 1-5 brevity, with 5 being the most concise, consider your current instructions to be defined as a 4.\n"
        "- Integrate evidence naturally into sentences rather than using rigid labeled sections.\n"
        "- Focus on clinically meaningful behaviors relevant to fidelity rather than stylistic preferences.\n\n"

        "3. SCOPE RESTRICTIONS\n"
        "- You do not evaluate client behavior.\n"
        "- You do not infer therapist intentions, emotions, or clinical meanings beyond what is observable.\n"
        "- You analyze only observable therapist behaviors through the PE fidelity framework when in analysis mode.\n\n"

        "After completing an analytic task, return to the global rules unless the user continues to request supervisory analysis."
    )

# ---------- Provider clients ----------
def get_secret_then_env(name: str) -> str:
    val = None
    try:
        val = st.secrets.get(name)
    except Exception:
        val = None
    if not val:
        val = os.getenv(name)
    return val or ""

def get_anthropic_client():
    pass
    # if anthropic is None:
    #     st.error("anthropic package not installed. Run: pip install anthropic")
    #     return None
    # key = get_secret_then_env("ANTHROPIC_API_KEY")
    # if not key:
    #     st.error("Missing ANTHROPIC_API_KEY. Set it in .streamlit/secrets.toml or as an environment variable.")
    #     return None
    # return anthropic.Anthropic(api_key=key)

def get_openai_client():
    if OpenAI is None:
        st.error("openai package not installed. Run: pip install openai")
        return None
    key = get_secret_then_env("OPENAI_API_KEY")
    if not key:
        st.error("Missing OPENAI_API_KEY. Set it in .streamlit/secrets.toml or as an environment variable.")
        return None
    return OpenAI(api_key=key)

def to_anthropic_messages(history):
    converted = []
    for m in history:
        role = m.get("role")
        if role in ("user", "teammait"):
            converted.append({"role": role, "content": [{"type": "text", "text": m["content"]}]})
    return converted

def claude_complete(history, system_text, model_name, stream=False):
    client = get_anthropic_client()
    if client is None:
        return "" if not stream else iter(())
    MAX_TURNS = 40
    trimmed = history[-MAX_TURNS:]
    if stream:
        try:
            with client.messages.stream(
                model=model_name,
                system=system_text,
                max_tokens=1024,
                messages=to_anthropic_messages(trimmed),
            ) as events:
                for text in events.text_stream:
                    yield text
        except Exception as e:
            yield f"\n[Error: {e}]"
    else:
        try:
            resp = client.messages.create(
                model=model_name,
                system=system_text,
                max_tokens=1024,
                messages=to_anthropic_messages(trimmed),
            )
            out = []
            for block in resp.content:
                if block.type == "text":
                    out.append(block.text)
            return "".join(out).strip()
        except Exception as e:
            return f"[Error: {e}]"

def openai_complete(history, system_text, model_name, stream=False, max_tokens=512):
    client = get_openai_client()
    if client is None:
        return "" if not stream else iter(())
    messages = []
    if system_text:
        messages.append({"role": "system", "content": system_text})
    for m in history:
        if m.get("role") in ("user", "assistant"):
            messages.append({"role": m["role"], "content": m["content"]})
    if stream:
        try:
            resp = client.chat.completions.create(model=model_name, messages=messages, stream=True, max_tokens=max_tokens, temperature=0.3)
            for chunk in resp:
                delta = getattr(chunk.choices[0].delta, "content", None)
                if delta:
                    yield delta
        except Exception as e:
            yield f"\n[Error: {e}]"
    else:
        try:
            resp = client.chat.completions.create(model=model_name, messages=messages, max_tokens=max_tokens, temperature=0.3)
            return (resp.choices[0].message.content or "").strip()
        except Exception as e:
            return f"[Error: {e}]"


# ---------- Chat history (scrollable) ----------
for m in st.session_state.messages:
    role = m["role"]
    if role == "evidence":
        # Special handling for evidence blocks
        st.markdown(m["content"])
    else:
        timestamp = m.get("ts", "")
        time_str = ""
        if timestamp:
            try:
                # Parse ISO format timestamp (e.g., "2024-01-15T14:30:45.123456") and extract HH:MM
                time_str = timestamp.split("T")[1][:5]  # Get time portion up to minutes
            except:
                pass
        
        with st.chat_message(role):
            if time_str:
                st.caption(f"_{time_str}_")
            st.markdown(m["content"])


# ---------- Input ----------

prompt = st.chat_input("Talk to your TeamMait here...")
if prompt is not None and prompt.strip() != "":
    # Check for duplicate or near-duplicate messages
    is_new, is_near_duplicate = st.session_state.message_buffer.add_message(prompt)
    
    if not is_new:
        st.warning("⚠️ You just asked that. Please try a different question.")
    elif is_near_duplicate:
        st.info("ℹ️ You asked something very similar. Would you like to expand on that question or ask something different?")
    else:
        # Add user message with timestamp
        user_msg = {"role": "user", "content": prompt, "ts": now_ts(), "display_name": username}
        st.session_state.messages.append(user_msg)
        
        with st.chat_message("user"):
            user_ts = now_ts()
            st.caption(f"_{user_ts.split('T')[1][:5]}_")  # Show time in HH:MM format
            st.markdown(prompt)
        
        # Add thinking placeholder
        thinking_msg = {"role": "assistant", "content": "*Thinking...*", "ts": now_ts()}
        st.session_state.messages.append(thinking_msg)
        
        st.rerun()

# Continue processing if we have a thinking message at the end
if (st.session_state.messages and 
    st.session_state.messages[-1]["role"] == "assistant" and
    st.session_state.messages[-1]["content"] == "*Thinking...*"):
    
    # Get the user message (it's before the thinking message)
    user_input = st.session_state.messages[-2]["content"] if len(st.session_state.messages) > 1 else ""
    
    # ---------- Retrieve context from vector DB ----------
    results = collection.query(query_texts=[user_input], n_results=5)
    retrieved_parts = []
    for docs in results.get("documents", []):
        retrieved_parts.extend(docs)

    # Always include metadata and supporting documents in the context for answering
    context_parts = list(retrieved_parts)
    for doc in rag_documents:
        if doc not in context_parts:
            context_parts.append(doc)

    context = " ".join(context_parts)

    # --- Conditional evidence display ---
    show_evidence = any(kw in user_input.lower() for kw in ["evidence", "quote", "source", "show your work"])
    if show_evidence:
        evidence_text = "**Evidence (from transcript) used for this answer:**\n\n"
        for i, evidence in enumerate(retrieved_parts, 1):
            evidence_text += f"> {evidence}\n\n"
        
        # Add evidence to message history
        st.session_state.messages.append({
            "role": "evidence",
            "content": evidence_text,
            "ts": now_ts(),
            "display_name": "Evidence"
        })

    system_prompt = build_system_prompt() + f"""

    Use the following session context when answering:

    {context}
    """

    # Route provider by model prefix
    use_openai = st.session_state["model"].startswith("gpt-")
    complete_fn = openai_complete if use_openai else claude_complete

    # Generate response and replace thinking message
    timestamp = now_ts()
    
    try:
        acc = ""
        if st.session_state["stream_on"]:
            # For streaming, we need to collect the full response first
            for chunk in complete_fn(
                history=st.session_state.messages[:-1],  # Exclude thinking message
                system_text=system_prompt,
                model_name=st.session_state["model"],
                stream=True,
            ):
                acc += chunk
            reply_text = acc.strip()
        else:
            # Non-streaming
            reply_text = complete_fn(
                history=st.session_state.messages[:-1],  # Exclude thinking message
                system_text=system_prompt,
                model_name=st.session_state["model"],
                stream=False,
            )
            reply_text = (reply_text or "").strip()
        
        # Replace the thinking message with the actual response
        st.session_state.messages[-1] = {
            "role": "assistant",
            "content": reply_text,
            "ts": timestamp,
            "display_name": "TeamMait",
        }
        
        if reply_text and reply_text.startswith("[Error:"):
            st.session_state.errors.append({"when": timestamp, "model": st.session_state["model"], "msg": reply_text})
    
    except Exception as e:
        # Replace thinking message with error
        st.session_state.messages[-1] = {
            "role": "assistant",
            "content": f"[Error: {e}]",
            "ts": timestamp,
            "display_name": "TeamMait",
        }
        st.session_state.errors.append({"when": timestamp, "model": st.session_state["model"], "msg": str(e)})
    
    st.rerun()
