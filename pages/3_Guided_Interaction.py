import streamlit as st
import sys
import os
from datetime import datetime
from enum import Enum
from dataclasses import dataclass
from typing import Optional, List, Dict
import json
import glob
import random

# Fix the import path
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)

# SQLite shim for Chroma
try:
    import pysqlite3
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    pass

import chromadb
from chromadb.utils import embedding_functions
from chromadb.config import Settings, DEFAULT_TENANT, DEFAULT_DATABASE

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from docx import Document
except ImportError:
    Document = None

# ==================== STATE MACHINE DEFINITION ====================

class ChatbotState(Enum):
    """Define all possible chatbot states."""
    IDLE = "idle"
    QUESTION_PRESENTATION = "question_presentation"
    OPEN_DISCUSSION = "open_discussion"
    SESSION_COMPLETE = "session_complete"

@dataclass
class StateTransition:
    """Represents a state transition with associated data."""
    next_state: ChatbotState
    bot_response: Optional[str] = None
    use_llm: bool = False
    show_buttons: bool = False
    show_feedback_buttons: bool = False
    api_context: Optional[str] = None

class GuidedInteractionStateMachine:
    """State machine for managing guided interaction flow."""
    
    def __init__(self, question_bank: List[Dict], openai_client):
        self.current_state = ChatbotState.IDLE
        self.questions_asked = []
        self.current_question = None
        self.current_response_type = None
        self.question_bank = question_bank
        self.openai_client = openai_client
        self.button_clicked = None
        self.pending_offer = None  # Track if we offered something specific
        self.last_bot_message = None  # Track last bot message for context
        
    def set_last_bot_message(self, message: str):
        """Store the last bot message for context."""
        self.last_bot_message = message
        
        # Check if bot just made an offer for analysis
        if "would you like me to provide an analysis" in message.lower() or \
           "would you like an analysis" in message.lower():
            # Extract what analysis was offered
            self.pending_offer = "analysis"
        else:
            self.pending_offer = None
    
    def transition(self, user_input: str) -> StateTransition:
        """Main transition logic - routes to state-specific handlers."""
        
        # PRIORITY 1: Check if user is accepting a pending offer
        if self.pending_offer == "analysis" and self._is_accepting_offer(user_input):
            return self._provide_promised_analysis(user_input)
        
        # PRIORITY 2: Check for next question request (but not if in middle of discussion)
        if self.current_state != ChatbotState.IDLE and \
           not self.pending_offer and \
           self._wants_next_question(user_input):
            return self._advance_to_next_question()
        
        # Route to state-specific handler
        if self.current_state == ChatbotState.IDLE:
            return self._handle_idle(user_input)
        elif self.current_state == ChatbotState.QUESTION_PRESENTATION:
            return self._handle_question_presentation(user_input)
        elif self.current_state == ChatbotState.OPEN_DISCUSSION:
            return self._handle_open_discussion(user_input)
        elif self.current_state == ChatbotState.SESSION_COMPLETE:
            return self._handle_complete(user_input)
        
        # Fallback
        return StateTransition(next_state=self.current_state)
    
    def _is_accepting_offer(self, user_input: str) -> bool:
        """Check if user is accepting an offered analysis."""
        acceptance_words = ["yes", "sure", "ok", "okay", "please", "go ahead", "yeah"]
        user_lower = user_input.lower().strip()
        return user_lower in acceptance_words or any(word in user_lower for word in acceptance_words)
    
    def _provide_promised_analysis(self, user_input: str) -> StateTransition:
        """Provide the analysis that was previously offered."""
        current_prompt = self.current_question or {}
        
        # Determine what analysis was offered based on last message
        analysis_topic = "the session"
        if self.last_bot_message:
            if "rapport and empathic connection" in self.last_bot_message.lower():
                analysis_topic = "rapport and empathic connection across the session"
            elif "therapeutic alliance" in self.last_bot_message.lower():
                analysis_topic = "therapeutic alliance in the session"
        
        analysis_context = f"""The user accepted your offer to provide an analysis about {analysis_topic}.

Your previous message offered: "{self.last_bot_message}"

The user said: "{user_input}"

Now provide the detailed, specific analysis that you offered. Focus on {analysis_topic} with:
- Specific evidence from the transcript with line number citations
- Clear observations about therapist skills and techniques
- Concrete examples from the conversation
- Professional assessment of effectiveness

Be direct and provide the substantive analysis you promised. Do NOT offer more analysis or ask more questions - just deliver what you promised."""

        # Clear the pending offer
        self.pending_offer = None
        
        return StateTransition(
            next_state=ChatbotState.OPEN_DISCUSSION,
            use_llm=True,
            api_context=analysis_context
        )
    
    # ... rest of the methods stay the same ...
    
    def _wants_next_question(self, user_input: str) -> bool:
        """Detect if user wants the next question using LLM."""
        # Don't trigger on simple acceptances
        user_lower = user_input.lower().strip()
        simple_acceptances = ["yes", "sure", "ok", "okay", "yeah", "yep"]
        if user_lower in simple_acceptances:
            return False
        
        if not self.openai_client:
            # Fallback to keyword matching
            keywords = [
                "next question", "next observation", "what's next", "another question",
                "move on", "next"
            ]
            return any(kw in user_input.lower() for kw in keywords)
        
        try:
            prompt = f"""Does this user response indicate they want to proceed to the next structured observation/question?

User said: "{user_input}"

Context: They are currently discussing a therapy transcript observation.

If they're just agreeing or saying "yes/sure/ok", that's NOT requesting the next question.
Only return "yes" if they're explicitly asking for the next question/observation.

Respond with ONLY "yes" or "no"."""

            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=5,
                temperature=0.1
            )
            return response.choices[0].message.content.strip().lower() == "yes"
        except:
            keywords = ["next question", "next observation", "what's next", "move on"]
            return any(kw in user_input.lower() for kw in keywords)
    
    def _advance_to_next_question(self) -> StateTransition:
        """Advance to the next question in the bank."""
        # Get next question that hasn't been asked
        remaining_questions = [q for q in self.question_bank if q not in self.questions_asked]
        
        if remaining_questions:
            next_question = remaining_questions[0]
            self.questions_asked.append(next_question)
            self.current_question = next_question
            
            formatted_response = f"{next_question.get('assertion', '')}\n\n{next_question.get('explanation', '')}\n\n{next_question.get('invitation', '')}"
            
            return StateTransition(
                next_state=ChatbotState.QUESTION_PRESENTATION,
                bot_response=formatted_response,
                show_buttons=True,
                show_feedback_buttons=True
            )
        else:
            # No more questions
            return StateTransition(
                next_state=ChatbotState.SESSION_COMPLETE,
                bot_response="All questions have been reviewed! Great work on finishing the guided interaction."
            )
    
    def _handle_idle(self, user_input: str) -> StateTransition:
        """Handle IDLE state - waiting for user to request first question."""
        if self._wants_next_question(user_input):
            return self._advance_to_next_question()
        
        # If user is just chatting (like saying "hi"), respond naturally with LLM
        casual_greetings = ["hi", "hello", "hey", "howdy", "greetings", "sup", "yo"]
        if user_input.lower().strip() in casual_greetings or len(user_input.strip()) < 15:
            return StateTransition(
                next_state=ChatbotState.IDLE,
                use_llm=True,  # Use LLM to respond naturally
                api_context=f"""The user said: "{user_input}"
                
    This appears to be casual conversation or a greeting. Respond warmly and naturally as a peer clinician would. 
    Remind them that you have observations to share when they're ready, but do it conversationally.
    Keep it brief and friendly."""
            )
        
        return StateTransition(
            next_state=ChatbotState.IDLE,
            bot_response="I'm ready to share my observations when you are. Just ask for the next question when you'd like to begin.",
            show_buttons=False,
            show_feedback_buttons=False
        )
    
    def _handle_question_presentation(self, user_input: str) -> StateTransition:
        """Handle user response to a presented question."""
        # Classify response type
        response_type = self._classify_response(user_input)
        
        if response_type == "disregard":
            return StateTransition(
                next_state=ChatbotState.OPEN_DISCUSSION,
                bot_response="Noted. Feel free to ask me anything else about this transcript, or request the next question when you're ready."
            )
        elif response_type == "accept":
            return StateTransition(
                next_state=ChatbotState.OPEN_DISCUSSION,
                bot_response="I'm glad you agree with that observation. Is there anything else you'd like to discuss about this topic, or are you ready to move on to my next observation?"
            )
        elif response_type in ["correct", "clarify"]:
            return StateTransition(
                next_state=ChatbotState.OPEN_DISCUSSION,
                bot_response="I appreciate your perspective. What specifically would you like to discuss or clarify about this observation?",
                use_llm=True
            )
        else:
            # Unclear response - use LLM to respond naturally
            return StateTransition(
                next_state=ChatbotState.QUESTION_PRESENTATION,
                use_llm=True,
                show_buttons=True,
                show_feedback_buttons=True
            )
    
    def _handle_open_discussion(self, user_input: str) -> StateTransition:
        """Handle user input during open discussion."""
        return StateTransition(
            next_state=ChatbotState.OPEN_DISCUSSION,
            use_llm=True
        )
    
    def _handle_complete(self, user_input: str) -> StateTransition:
        """Handle user input when session is complete."""
        return StateTransition(
            next_state=ChatbotState.SESSION_COMPLETE,
            bot_response="The guided interaction has been completed. Thank you for your participation!"
        )
    
    def _classify_response(self, user_input: str) -> str:
        """Classify user response to a question."""
        if not self.openai_client:
            return "unclear"
        
        try:
            prompt = f"""Classify this response to a clinical observation:

User response: "{user_input}"

Categories:
- "accept" - agrees with observation
- "correct" - disagrees or wants to correct
- "clarify" - asks for clarification
- "disregard" - wants to move on/dismiss
- "unclear" - doesn't fit above

Respond with ONLY the category name."""

            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=10,
                temperature=0.1
            )
            return response.choices[0].message.content.strip().lower()
        except:
            return "unclear"

# ==================== HELPER FUNCTIONS ====================

def now_ts() -> str:
    """Get current timestamp."""
    return datetime.now().strftime("%H:%M:%S")

def get_secret_then_env(name: str) -> str:
    """Get secret from Streamlit secrets or environment variable."""
    val = None
    try:
        val = st.secrets.get(name)
    except Exception:
        val = None
    if not val:
        val = os.getenv(name)
    return val or ""

def get_openai_client():
    """Initialize OpenAI client."""
    if OpenAI is None:
        st.error("openai package not installed. Run: pip install openai")
        return None
    key = get_secret_then_env("OPENAI_API_KEY")
    if not key:
        st.error("Missing OPENAI_API_KEY.")
        return None
    return OpenAI(api_key=key)

def build_system_prompt() -> str:
    """Build system prompt for guided mode."""
    return (
        "You are TeamMait, a peer-support assistant for expert clinicians reviewing "
        "therapist performance in a transcript. Your scope is limited strictly to "
        "analyzing the therapist's observable skills in the transcript. "
        "Anchor every claim to the transcript (and provided docs). When citing the transcript, "
        "use the line numbers provided in the format [Line X]. If uncertain, say so briefly. "
        "Be succinct and academically neutral; do not use emojis. "
        "Never invent facts. Always cite specific line references when making claims about the transcript. "
        
        "\n\nYou are operating in GUIDED REVIEW mode. You present observations about the transcript "
        "and engage in natural discussion with the clinician. When they respond to your observations, "
        "engage conversationally and provide thoughtful analysis. "
        
        "\n\nCONVERSATION GUIDELINES:"
        "\n1. AVOID REPETITIVE PHRASES: NEVER start with 'It sounds like...' or 'It seems like...'. "
        "\n2. ANSWER DIRECT QUESTIONS: When users ask questions, answer them directly with substantive guidance."
        "\n3. SUBSTANTIVE ENGAGEMENT: Provide analysis and insights, not just questions."
        "\n4. Mimic the tone and style of the collaborative peer clinician you are conversing with."
    )

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text content from a DOCX file."""
    if Document is None:
        return f"[Error: python-docx not installed]"
    
    try:
        doc = Document(docx_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"[Error: {str(e)}]"

def openai_complete(history, system_text, client, model_name="gpt-4o-mini", stream=False, max_tokens=512):
    """Complete a chat using OpenAI API."""
    if client is None:
        return "" if not stream else iter(())
    
    messages = []
    if system_text:
        messages.append({"role": "system", "content": system_text})
    
    for m in history:
        if m.get("role") in ("user", "assistant"):
            messages.append({"role": m["role"], "content": m["content"]})
    
    if stream:
        try:
            resp = client.chat.completions.create(
                model=model_name,
                messages=messages,
                stream=True,
                max_tokens=max_tokens,
                temperature=0.3
            )
            for chunk in resp:
                delta = getattr(chunk.choices[0].delta, "content", None)
                if delta:
                    yield delta
        except Exception as e:
            yield f"\n[Error: {e}]"
    else:
        try:
            resp = client.chat.completions.create(
                model=model_name,
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.3
            )
            return (resp.choices[0].message.content or "").strip()
        except Exception as e:
            return f"[Error: {e}]"

# ==================== RAG SETUP ====================

@st.cache_resource(show_spinner=False)
def initialize_chroma():
    """Initialize ChromaDB client and collection."""
    embed_model = "sentence-transformers/all-MiniLM-L6-v2"
    embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=embed_model
    )
    
    chroma_client = chromadb.PersistentClient(
        path="./rag_store",
        settings=Settings(),
        tenant=DEFAULT_TENANT,
        database=DEFAULT_DATABASE,
    )
    
    collection = chroma_client.get_or_create_collection(
        "therapy",
        embedding_function=embedding_fn
    )
    
    return chroma_client, collection

@st.cache_resource(show_spinner=False)
def load_rag_documents():
    """Load all RAG documents and seed ChromaDB collection."""
    _, collection = initialize_chroma()
    
    doc_folder = "doc/RAG"
    supporting_folder = os.path.join(doc_folder, "supporting_documents")
    documents = []
    ids = []
    
    # Load main reference conversation with line numbers
    ref_path = os.path.join(doc_folder, "116_P8_conversation.json")
    if os.path.exists(ref_path):
        with open(ref_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict) and "full_conversation" in data:
                for i, turn in enumerate(data["full_conversation"]):
                    line_num = i + 1
                    numbered_turn = f"[Line {line_num}] {turn}"
                    documents.append(numbered_turn)
                    ids.append(f"ref_{i}")
    
    # Load supporting documents
    for txt_path in glob.glob(os.path.join(supporting_folder, "*.txt")):
        with open(txt_path, "r", encoding="utf-8") as f:
            documents.append(f.read())
            ids.append(f"supp_txt_{os.path.basename(txt_path)}")
    
    for json_path in glob.glob(os.path.join(supporting_folder, "*.json")):
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, list):
                for i, item in enumerate(data):
                    documents.append(str(item))
                    ids.append(f"supp_json_{os.path.basename(json_path)}_{i}")
    
    for docx_path in glob.glob(os.path.join(supporting_folder, "*.docx")):
        content = extract_text_from_docx(docx_path)
        if content and not content.startswith("[Error"):
            documents.append(content)
            ids.append(f"supp_docx_{os.path.basename(docx_path)}")
    
    # Seed collection if empty
    if collection.count() == 0 and documents:
        collection.add(documents=documents, ids=ids)
    
    return documents

def retrieve_context(query: str, n_results: int = 5) -> str:
    """Retrieve relevant context from ChromaDB."""
    _, collection = initialize_chroma()
    results = collection.query(query_texts=[query], n_results=n_results)
    
    retrieved_parts = []
    for docs in results.get("documents", []):
        retrieved_parts.extend(docs)
    
    return " ".join(retrieved_parts)

def load_reference_conversation():
    """Load the reference conversation for display."""
    ref_path = os.path.join("doc/RAG", "116_P8_conversation.json")
    if os.path.exists(ref_path):
        with open(ref_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict) and "full_conversation" in data:
                return data["full_conversation"]
    return []

def load_question_bank():
    """Load the question bank."""
    with open("doc/interaction_prompts/interaction_prompts.json", "r") as f:
        data = json.load(f)
    return data.get("feedback_items", [])

# ==================== STREAMLIT APP ====================

st.set_page_config(page_title="Guided Interaction", page_icon="💬", layout="wide")

# Check login
if "user_info" not in st.session_state:
    st.warning("Please log in first.")
    st.switch_page("Home.py")
    st.stop()

username = st.session_state["user_info"]["username"]

# Initialize OpenAI client
client = get_openai_client()

# Initialize state machine
if "state_machine" not in st.session_state:
    question_bank = load_question_bank()
    st.session_state.state_machine = GuidedInteractionStateMachine(question_bank, client)

# Initialize message history
if "guided_messages" not in st.session_state:
    st.session_state.guided_messages = [
        {
            "role": "assistant",
            "content": "Hi, my name is TeamMait. Feel free to ask me any questions about the referenced session transcript. It can be found in the left side panel.\n\nI've made a few observations about the session that we can discuss together. When you're ready, just ask me for the 'next question' and I'll share one with you.",
            "ts": now_ts()
        }
    ]

# Load RAG documents
load_rag_documents()

# ==================== SIDEBAR ====================

with st.sidebar:
    st.markdown(f"**Username:** {username}")
    
    # Completion status
    if "completion_status" not in st.session_state:
        st.session_state["completion_status"] = {}
    
    persistent_value = st.session_state["completion_status"].get("guided_interaction", False)
    st.session_state["include_guided_interaction"] = persistent_value

    def _on_include_guided_change():
        current_value = st.session_state.get("include_guided_interaction", False)
        st.session_state["completion_status"]["guided_interaction"] = current_value

    st.checkbox("Check this when done", key="include_guided_interaction", on_change=_on_include_guided_change)
    
    # Progress tracker
    st.markdown("### Progress")
    questions_asked = len(st.session_state.state_machine.questions_asked)
    st.progress(min(questions_asked / 4, 1.0))
    st.caption(f"{questions_asked} / 4 questions reviewed")
    
    # Show reference conversation
    with st.expander("Show Reference Conversation", expanded=True):
        ref_conversation = load_reference_conversation()
        if ref_conversation:
            for i, turn in enumerate(ref_conversation):
                line_num = i + 1
                is_client = turn.strip().startswith("Client: ")
                if is_client:
                    st.markdown(f"""
                    <div style="text-align: right; margin-left: 0%; padding: 10px; border-radius: 10px;">
                    <small style="color: #888; font-size: 0.8em;">[Line {line_num}]</small><br>
                    {turn}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"<div style='font-weight:600; font-size:1.08em; margin: 10px 0;'><small style='color: #888; font-size: 0.8em; font-weight: normal;'>[Line {line_num}]</small><br><em>{turn}</em></div>", unsafe_allow_html=True)

# ==================== MAIN CONTENT ====================

st.title("Guided Interaction")
st.markdown("<p style='font-size:12px;color:#e11d48;margin-top:6px;'><strong>Privacy Reminder:</strong> Please do not include any identifying information in your messages.</p>", unsafe_allow_html=True)
st.markdown("<p style='font-size:12px;color:#6b7280;margin-bottom:6px;'>Disclaimer: TeamMait may be incorrect or incomplete. Please verify information.</p>", unsafe_allow_html=True)

# Display chat history
for m in st.session_state.guided_messages:
    role = m["role"]
    role_label = "TeamMait" if role == "assistant" else "User"
    
    with st.chat_message(role):
        if "ts" in m:
            st.markdown(f"**{role_label}** • <small style='color: #888; font-size: 0.8em;'>*{m['ts']}*</small>", unsafe_allow_html=True)
        else:
            st.markdown(f"**{role_label}**")
        st.markdown(m["content"])

# Response buttons
st.write(f"DEBUG: Current state is {st.session_state.state_machine.current_state}")  # Debug line

def handle_button_click(response_type: str):
    """Handle quick response button clicks."""
    st.session_state.state_machine.button_clicked = response_type
    st.rerun()

if st.session_state.state_machine.current_state == ChatbotState.QUESTION_PRESENTATION:
    st.write(f"DEBUG: In QUESTION_PRESENTATION state") # Debug line
    st.markdown("##### Grade Response (optional):")
    cols = st.columns(2)
    
    button_clicked = st.session_state.state_machine.button_clicked
    
    with cols[0]:
        button_text = "👍 ✓" if button_clicked == "accept" else "👍"
        st.button(button_text, on_click=lambda: handle_button_click("accept"), use_container_width=True)
    with cols[1]:
        button_text = "👎 ✓" if button_clicked == "disregard" else "👎"
        st.button(button_text, on_click=lambda: handle_button_click("disregard"), use_container_width=True)

# Text input
if st.session_state.state_machine.current_state != ChatbotState.SESSION_COMPLETE:
    prompt = st.chat_input("Type your response or elaboration...")
    
    if prompt:
        # Add user message
        user_msg = {
            "role": "user",
            "content": prompt,
            "ts": now_ts()
        }
        st.session_state.guided_messages.append(user_msg)
        
        with st.chat_message("user"):
            st.markdown(f"**User** • <small style='color: #888; font-size: 0.8em;'>*{user_msg['ts']}*</small>", unsafe_allow_html=True)
            st.markdown(prompt)
        
        # Process state transition
        transition = st.session_state.state_machine.transition(prompt)
        
        # Generate response
        if transition.use_llm:
            # Use LLM to generate response
            if transition.api_context:
                context = transition.api_context
            else:
                context = retrieve_context(prompt)
            
            system_prompt = build_system_prompt() + f"\n\nContext:\n{context}"
            
            with st.chat_message("assistant"):
                timestamp = now_ts()
                st.markdown(f"**TeamMait** • <small style='color: #888; font-size: 0.8em;'>*{timestamp}*</small>", unsafe_allow_html=True)
                
                placeholder = st.empty()
                acc = ""
                for chunk in openai_complete(
                    history=[{"role": "user", "content": prompt}],
                    system_text=system_prompt,
                    client=client,
                    stream=True,
                    max_tokens=512
                ):
                    acc += chunk
                    placeholder.markdown(acc)
                
                bot_msg = {
                    "role": "assistant",
                    "content": acc.strip(),
                    "ts": timestamp
                }
                st.session_state.guided_messages.append(bot_msg)
                # ADD THIS LINE HERE:
                st.session_state.state_machine.set_last_bot_message(acc.strip())         
        elif transition.bot_response:
            timestamp = now_ts()
            with st.chat_message("assistant"):
                st.markdown(f"**TeamMait** • <small style='color: #888; font-size: 0.8em;'>*{timestamp}*</small>", unsafe_allow_html=True)
                st.markdown(transition.bot_response)
                
            bot_msg = {
                "role": "assistant",
                "content": transition.bot_response,
                "ts": timestamp
            }
            st.session_state.guided_messages.append(bot_msg)
            # ADD THIS LINE HERE:
            st.session_state.state_machine.set_last_bot_message(transition.bot_response)
        
        # Update state machine state
        st.session_state.state_machine.current_state = transition.next_state
        
        # Clear button state
        st.session_state.state_machine.button_clicked = None
        
        st.rerun()

else:
    st.success("Guided interaction session complete!")
    if st.button("Start New Session"):
        question_bank = load_question_bank()
        st.session_state.state_machine = GuidedInteractionStateMachine(question_bank, client)
        st.session_state.guided_messages = [{
                "role": "assistant",
                "content": "Hi, my name is TeamMait. Feel free to ask me any questions about the referenced session transcript. It can be found in the left side panel.\n\nI've made a few observations about the session that we can discuss together. When you're ready, just ask me for the 'next question' and I'll share one with you.",
                "ts": now_ts()
            }
        ]
        st.rerun()