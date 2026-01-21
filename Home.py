"""
TeamMait - AI-Powered Therapy Transcript Review Assistant
Main chat interface for clinicians to review and discuss therapy session transcripts.
"""

import streamlit as st
import json
import uuid
from datetime import datetime
from textwrap import dedent
from urllib.parse import quote
import os
import sys
import glob

# ---------- SQLite shim for Chroma ----------
try:
    import pysqlite3
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    pass

import chromadb
from chromadb.utils import embedding_functions
from chromadb.config import Settings, DEFAULT_TENANT, DEFAULT_DATABASE

# Optional SDKs
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from docx import Document
except ImportError:
    Document = None

from utils.session_manager import SessionManager
from utils.input_parser import MessageBuffer

# ---------- Page Config ----------
st.set_page_config(
    page_title="TeamMait - Therapy Transcript Review",
    layout="wide"
)

# ---------- Simple avatars (SVG data URIs) ----------
def svg_data_uri(svg: str) -> str:
    return "data:image/svg+xml;utf8," + quote(svg)

USER_SVG = svg_data_uri(
    dedent("""
    <svg xmlns='http://www.w3.org/2000/svg' width='64' height='64'>
      <defs>
        <linearGradient id='g' x1='0' x2='1' y1='0' y2='1'>
          <stop offset='0%' stop-color='#1f2937'/>
          <stop offset='100%' stop-color='#0b1220'/>
        </linearGradient>
      </defs>
      <circle cx='32' cy='32' r='32' fill='url(#g)'/>
      <text x='50%' y='54%' text-anchor='middle' font-family='Inter,Arial' font-size='24' fill='#e5e7eb' font-weight='700'>U</text>
    </svg>
    """).strip()
)

BOT_SVG = svg_data_uri(
    dedent("""
    <svg xmlns='http://www.w3.org/2000/svg' width='64' height='64'>
      <circle cx='32' cy='32' r='32' fill='#111827'/>
      <rect x='18' y='20' width='28' height='22' rx='6' fill='#1f2937' stroke='#374151' stroke-width='2'/>
      <circle cx='26' cy='31' r='4' fill='#93c5fd'/>
      <circle cx='38' cy='31' r='4' fill='#93c5fd'/>
      <rect x='28' y='42' width='8' height='6' rx='3' fill='#4b5563'/>
      <rect x='30' y='12' width='4' height='8' rx='2' fill='#6b7280'/>
    </svg>
    """).strip()
)


# ---------- Helpers ----------
def now_ts() -> str:
    return datetime.now().isoformat()


def get_secret_then_env(name: str) -> str:
    val = None
    try:
        val = st.secrets.get(name)
    except Exception:
        val = None
    if not val:
        val = os.getenv(name)
    return val or ""


# ---------- Login Dialog ----------
def load_valid_users():
    """Load valid users from Streamlit secrets"""
    try:
        users = st.secrets.get("credentials", {}).get("users", [])
        if not users:
            return []
        return users
    except Exception:
        return []


def validate_login(username, password):
    """Validate username and password against secrets"""
    # Allow demo mode
    if username == "demo" and password == "demo":
        return True
    
    valid_users = load_valid_users()
    for user in valid_users:
        if user.get("username") == username and user.get("password") == password:
            return True
    return False


@st.dialog("Login", width="small")
def get_user_details():
    st.markdown("### Welcome to TeamMait")
    st.markdown("Enter your credentials to continue, or use `demo`/`demo` for a demo.")
    
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    submit = st.button("Login", type="primary", use_container_width=True)
    
    if submit:
        if not username or not password:
            st.error("Username and password are required.")
            return
        
        if not validate_login(username, password):
            st.error("Invalid credentials. Please try again.")
            return
        
        # Check if demo user needs API key
        if username == "demo" and password == "demo":
            st.session_state.is_demo_user = True
        else:
            st.session_state.is_demo_user = False
        
        st.session_state.user_info = {
            "username": username,
            "logged_in_at": datetime.now().isoformat()
        }
        st.session_state["username"] = username
        st.rerun()


# Check login
if "user_info" not in st.session_state:
    get_user_details()
    st.stop()

username = st.session_state["username"]

# Handle demo user API key
if st.session_state.get("is_demo_user") and "demo_api_key" not in st.session_state:
    st.info("**Demo Mode**: You'll need to provide your own OpenAI API key.")
    with st.form("api_key_form"):
        api_key_input = st.text_input(
            "Enter your OpenAI API Key",
            type="password",
            help="Your key will only be used for this session and will not be saved."
        )
        if st.form_submit_button("Save API Key", type="primary"):
            if not api_key_input or not api_key_input.startswith("sk-"):
                st.error("Please enter a valid OpenAI API key (starts with 'sk-').")
            else:
                st.session_state.demo_api_key = api_key_input
                st.rerun()
    st.stop()


# ---------- Initialize Session State ----------
if "main_session_id" not in st.session_state:
    st.session_state.main_session_id = str(uuid.uuid4())

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "assistant",
            "content": "Hi! I'm TeamMait, your peer-support assistant for reviewing therapy session transcripts. "
                      "You can find the session transcript in the sidebar. Feel free to ask me any questions about "
                      "the therapist's techniques, the session dynamics, or anything else you'd like to discuss.",
            "ts": now_ts(),
        }
    ]

if "message_buffer" not in st.session_state:
    st.session_state.message_buffer = MessageBuffer()


# ---------- OpenAI Client ----------
def get_openai_client():
    if OpenAI is None:
        st.error("OpenAI package not installed. Run: pip install openai")
        return None
    
    # Check for demo user API key first
    if st.session_state.get("is_demo_user") and st.session_state.get("demo_api_key"):
        key = st.session_state.get("demo_api_key")
    else:
        key = get_secret_then_env("OPENAI_API_KEY")
    
    if not key:
        st.error("Missing OPENAI_API_KEY.")
        return None
    return OpenAI(api_key=key)


# ---------- System Prompt ----------
def build_system_prompt() -> str:
    """Build system prompt for open chat mode."""
    return (
        "You are TeamMait, a peer-support assistant for expert clinicians reviewing therapist performance in a transcript. "
        "Your responses must follow two behavioral modes: global rules (always active) and analysis mode (only when the user requests supervisory analysis).\n\n"

        "1. GLOBAL RULES (ALWAYS ACTIVE)\n"
        "- Never fabricate transcript content, facts, or therapist intentions.\n"
        "- Do not infer internal states, emotions, or off-transcript behavior.\n"
        "- Respond concisely, professionally, and only to what the user asked.\n"
        "- Use a natural, peer-like supervisory tone; avoid rigid sections or templates unless the user requests structure.\n"
        "- Treat each user message independently unless the user explicitly references earlier turns.\n"
        "- If the user gives a dismissive acknowledgment (e.g., 'ok', 'thanks', 'got it'), briefly acknowledge and ask whether they want to continue or move on.\n"
        "- If the user expresses confusion, provide a simpler, more direct explanation of your prior point.\n"
        "- Do not offer unsolicited elaboration or additional insights outside analytic tasks.\n"
        "- Do NOT ask if the user wants you to analyze or offer to analyze. Simply respond to their question or request directly.\n"
        "- NEVER end messages with offers like 'Would you like me to...', 'Should I...', etc.\n\n"

        "2. ANALYSIS MODE (ONLY WHEN USER REQUESTS SUPERVISORY ANALYSIS)\n"
        "Enter analysis mode only when the user asks you to analyze therapist behavior, evaluate fidelity, generate observations, or provide supervision-like feedback.\n\n"

        "When in analysis mode:\n"
        "- Cite transcript lines in the format [Line X] when providing evidence.\n"
        "- Base all claims on observable behavior only.\n"
        "- Use PE fidelity criteria as the interpretive framework.\n"
        "- Use calibrated language such as 'appears', 'may indicate', or 'based on [Line X–Y]'.\n"
        "- Frame feedback as observations or suggestions, not directives.\n"
        "- Present reasoning in a way that allows the clinician to agree, disagree, or reinterpret.\n"
        "- Do not generalize beyond this specific transcript or session.\n\n"

        "Format in Analysis Mode:\n"
        "- Provide ideally 3, but no more than 5, bullet points to answer a given query.\n"
        "- One sentence per bullet; limit each bullet to about 10 words or 75 characters.\n"
        "- Prioritize clarity and brevity; avoid redundancy.\n\n"

        "3. SCOPE RESTRICTIONS\n"
        "- You do not evaluate client behavior.\n"
        "- You do not infer therapist intentions, emotions, or clinical meanings beyond what is observable.\n"
        "- You analyze only observable therapist behaviors through the PE fidelity framework when in analysis mode.\n\n"

        "After completing an analytic task, return to the global rules unless the user continues to request supervisory analysis."
    )


# ---------- ChromaDB / RAG Setup ----------
@st.cache_resource(show_spinner=False)
def initialize_chroma():
    embed_model = "sentence-transformers/all-MiniLM-L6-v2"
    embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(model_name=embed_model)
    
    chroma_client = chromadb.PersistentClient(
        path="./rag_store",
        settings=Settings(),
        tenant=DEFAULT_TENANT,
        database=DEFAULT_DATABASE,
    )
    collection = chroma_client.get_or_create_collection("therapy", embedding_function=embedding_fn)
    return chroma_client, collection


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text content from a DOCX file."""
    if Document is None:
        return f"[Error: python-docx not installed]"
    
    try:
        doc = Document(docx_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"[Error reading {os.path.basename(docx_path)}: {str(e)}]"


@st.cache_resource(show_spinner=False)
def load_rag_documents():
    """Load RAG documents from doc/RAG folder."""
    _, collection = initialize_chroma()
    doc_folder = "doc/RAG"
    supporting_folder = os.path.join(doc_folder, "supporting_documents")
    
    documents = []
    ids = []
    
    # Load reference conversation
    ref_path = os.path.join(doc_folder, "116_P8_conversation.json")
    if os.path.exists(ref_path):
        with open(ref_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict) and "full_conversation" in data:
                for i, turn in enumerate(data["full_conversation"]):
                    line_num = i + 1
                    numbered_turn = f"[Line {line_num}] {turn}"
                    documents.append(numbered_turn)
                    ids.append(f"ref_{i}")
    
    # Load supporting documents
    if os.path.exists(supporting_folder):
        for json_path in glob.glob(os.path.join(supporting_folder, "*.json")):
            try:
                with open(json_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, dict):
                        documents.append(str(data))
                        ids.append(f"supp_{os.path.basename(json_path)}")
            except Exception:
                pass
        
        for docx_path in glob.glob(os.path.join(supporting_folder, "*.docx")):
            content = extract_text_from_docx(docx_path)
            if content and not content.startswith("[Error"):
                documents.append(content)
                ids.append(f"supp_{os.path.basename(docx_path)}")
    
    # Seed collection if empty
    if collection.count() == 0 and documents:
        collection.add(documents=documents, ids=ids)
    
    return documents


# Load RAG docs
rag_documents = load_rag_documents()


def retrieve_context(query: str, n_results: int = 5) -> str:
    """Retrieve relevant context from ChromaDB."""
    try:
        _, collection = initialize_chroma()
        results = collection.query(query_texts=[query], n_results=n_results)
        retrieved_parts = []
        for docs in results.get("documents", []):
            retrieved_parts.extend(docs)
        return " ".join(retrieved_parts)
    except Exception:
        return ""


# ---------- OpenAI Completion ----------
def openai_complete(history, system_text, model_name="gpt-4o-mini", stream=True, max_tokens=512):
    client = get_openai_client()
    if client is None:
        return "" if not stream else iter(())
    
    messages = []
    if system_text:
        messages.append({"role": "system", "content": system_text})
    for m in history:
        if m.get("role") in ("user", "assistant"):
            messages.append({"role": m["role"], "content": m["content"]})
    
    if stream:
        try:
            resp = client.chat.completions.create(
                model=model_name, messages=messages, stream=True, 
                max_tokens=max_tokens, temperature=0.3
            )
            for chunk in resp:
                delta = getattr(chunk.choices[0].delta, "content", None)
                if delta:
                    yield delta
        except Exception as e:
            yield f"\n[Error: {e}]"
    else:
        try:
            resp = client.chat.completions.create(
                model=model_name, messages=messages, 
                max_tokens=max_tokens, temperature=0.3
            )
            return (resp.choices[0].message.content or "").strip()
        except Exception as e:
            return f"[Error: {e}]"


# ---------- Load Reference Conversation ----------
def load_reference_conversation():
    ref_path = "doc/RAG/116_P8_conversation.json"
    if os.path.exists(ref_path):
        with open(ref_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, dict) and "full_conversation" in data:
                return data["full_conversation"]
    return []


# ---------- Sidebar ----------
with st.sidebar:
    st.markdown(f"**User:** {username}")
    
    if st.session_state.get("is_demo_user"):
        with st.expander("API Key Settings"):
            st.write("Your OpenAI API key is set for this session.")
            if st.button("Reset API Key"):
                del st.session_state.demo_api_key
                st.rerun()
    
    if st.button("Clear Chat", type="secondary"):
        st.session_state.messages = [st.session_state.messages[0]]  # Keep welcome message
        st.rerun()
    
    st.divider()
    
    # Reference Conversation
    with st.expander("Session Transcript", expanded=True):
        ref_conversation = load_reference_conversation()
        if ref_conversation:
            for i, turn in enumerate(ref_conversation):
                line_num = i + 1
                is_client = turn.strip().startswith("Client:")
                speaker = "Client:" if is_client else "Therapist:"
                content = turn.replace(f"{speaker} ", "", 1).strip()
                
                if is_client:
                    st.markdown(f"""
                    <div style="text-align: right; margin-left: 0%; padding: 8px; border-radius: 8px; background: rgba(100,100,100,0.1); margin-bottom: 8px;">
                    <small style="color: #888; font-size: 0.75em;">[Line {line_num}] {speaker}</small><br>
                    <span style="font-size: 0.9em;">{content}</span>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div style="padding: 8px; margin-bottom: 8px;">
                    <small style="color: #888; font-size: 0.75em;">[Line {line_num}] {speaker}</small><br>
                    <em style="font-size: 0.9em;">{content}</em>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("No transcript loaded.")


# ---------- Main Content ----------
st.title("TeamMait")
st.markdown(
    "<p style='font-size:14px;color:#6b7280;margin-bottom:16px;'>"
    "Your AI peer-support assistant for reviewing therapy session transcripts. "
    "Ask questions about the session, therapist techniques, or request analysis."
    "</p>",
    unsafe_allow_html=True
)

# Chat display
for m in st.session_state.messages:
    role = m["role"]
    timestamp = m.get("ts", "")
    time_str = ""
    if timestamp:
        try:
            time_str = timestamp.split("T")[1][:5]
        except:
            pass
    
    with st.chat_message(role):
        if time_str:
            st.caption(f"_{time_str}_")
        st.markdown(m["content"])


# Chat input
prompt = st.chat_input("Ask TeamMait about the session...")

if prompt and prompt.strip():
    # Check for duplicates
    result = st.session_state.message_buffer.add_message(prompt)
    if isinstance(result, tuple):
        is_new, is_near_duplicate = result
    else:
        is_new, is_near_duplicate = True, False
    
    if not is_new:
        st.warning("⚠️ You just asked that. Please try a different question.")
    elif is_near_duplicate:
        st.info("You asked something very similar. Would you like to expand on that question?")
    else:
        # Add user message
        user_msg = {"role": "user", "content": prompt, "ts": now_ts()}
        st.session_state.messages.append(user_msg)
        
        with st.chat_message("user"):
            st.caption(f"_{now_ts().split('T')[1][:5]}_")
            st.markdown(prompt)
        
        # Generate response
        context = retrieve_context(prompt)
        system_prompt = build_system_prompt() + f"\n\nSession context:\n{context}"
        
        with st.chat_message("assistant"):
            placeholder = st.empty()
            placeholder.markdown("*Thinking...*")
            
            acc = ""
            first_chunk = True
            for chunk in openai_complete(
                history=st.session_state.messages,
                system_text=system_prompt,
                stream=True,
            ):
                if first_chunk:
                    placeholder.empty()
                    first_chunk = False
                acc += chunk
                placeholder.markdown(acc)
            
            # Save response
            st.session_state.messages.append({
                "role": "assistant",
                "content": acc.strip(),
                "ts": now_ts(),
            })
